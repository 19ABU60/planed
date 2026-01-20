from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, status, Query
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from fastapi.responses import StreamingResponse
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta, date
import jwt
import bcrypt
from io import BytesIO
import httpx
import asyncio

# Document exports
from docx import Document
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Set database for services module
from services.auth import set_database
set_database(db)

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'planed-secret-key-2025')
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 24

# Invitation Code for Registration
INVITATION_CODE = os.environ.get('INVITATION_CODE', 'LASP2026')

# Create the main app
app = FastAPI(title="PlanEd API", version="2.0.0")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============== GERMAN SCHOOL HOLIDAYS DATA ==============
GERMAN_HOLIDAYS_2025_2026 = {
    "bayern": [
        {"name": "Herbstferien 2025", "start": "2025-10-27", "end": "2025-10-31"},
        {"name": "Weihnachtsferien 2025/26", "start": "2025-12-22", "end": "2026-01-05"},
        {"name": "Winterferien 2026", "start": "2026-02-16", "end": "2026-02-20"},
        {"name": "Osterferien 2026", "start": "2026-03-30", "end": "2026-04-10"},
        {"name": "Pfingstferien 2026", "start": "2026-05-26", "end": "2026-06-05"},
        {"name": "Sommerferien 2026", "start": "2026-07-27", "end": "2026-09-07"},
    ],
    "nrw": [
        {"name": "Herbstferien 2025", "start": "2025-10-13", "end": "2025-10-25"},
        {"name": "Weihnachtsferien 2025/26", "start": "2025-12-22", "end": "2026-01-06"},
        {"name": "Osterferien 2026", "start": "2026-03-30", "end": "2026-04-11"},
        {"name": "Pfingstferien 2026", "start": "2026-05-26", "end": "2026-05-26"},
        {"name": "Sommerferien 2026", "start": "2026-06-29", "end": "2026-08-11"},
    ],
    "berlin": [
        {"name": "Herbstferien 2025", "start": "2025-10-20", "end": "2025-11-01"},
        {"name": "Weihnachtsferien 2025/26", "start": "2025-12-22", "end": "2026-01-02"},
        {"name": "Winterferien 2026", "start": "2026-02-02", "end": "2026-02-07"},
        {"name": "Osterferien 2026", "start": "2026-03-30", "end": "2026-04-10"},
        {"name": "Pfingstferien 2026", "start": "2026-05-15", "end": "2026-05-15"},
        {"name": "Sommerferien 2026", "start": "2026-07-09", "end": "2026-08-21"},
    ],
    "baden-wuerttemberg": [
        {"name": "Herbstferien 2025", "start": "2025-10-27", "end": "2025-10-30"},
        {"name": "Weihnachtsferien 2025/26", "start": "2025-12-22", "end": "2026-01-05"},
        {"name": "Osterferien 2026", "start": "2026-04-06", "end": "2026-04-17"},
        {"name": "Pfingstferien 2026", "start": "2026-05-26", "end": "2026-06-06"},
        {"name": "Sommerferien 2026", "start": "2026-07-30", "end": "2026-09-12"},
    ],
    "hessen": [
        {"name": "Herbstferien 2025", "start": "2025-10-06", "end": "2025-10-18"},
        {"name": "Weihnachtsferien 2025/26", "start": "2025-12-22", "end": "2026-01-10"},
        {"name": "Osterferien 2026", "start": "2026-04-06", "end": "2026-04-18"},
        {"name": "Sommerferien 2026", "start": "2026-07-06", "end": "2026-08-14"},
    ],
    "sachsen": [
        {"name": "Herbstferien 2025", "start": "2025-10-20", "end": "2025-11-01"},
        {"name": "Weihnachtsferien 2025/26", "start": "2025-12-22", "end": "2026-01-03"},
        {"name": "Winterferien 2026", "start": "2026-02-09", "end": "2026-02-21"},
        {"name": "Osterferien 2026", "start": "2026-04-03", "end": "2026-04-11"},
        {"name": "Pfingstferien 2026", "start": "2026-05-15", "end": "2026-05-15"},
        {"name": "Sommerferien 2026", "start": "2026-06-27", "end": "2026-08-08"},
    ],
    "niedersachsen": [
        {"name": "Herbstferien 2025", "start": "2025-10-20", "end": "2025-10-31"},
        {"name": "Weihnachtsferien 2025/26", "start": "2025-12-22", "end": "2026-01-05"},
        {"name": "Osterferien 2026", "start": "2026-03-23", "end": "2026-04-04"},
        {"name": "Pfingstferien 2026", "start": "2026-05-22", "end": "2026-05-22"},
        {"name": "Sommerferien 2026", "start": "2026-07-16", "end": "2026-08-26"},
    ],
    "hamburg": [
        {"name": "Herbstferien 2025", "start": "2025-10-20", "end": "2025-10-31"},
        {"name": "Weihnachtsferien 2025/26", "start": "2025-12-22", "end": "2026-01-02"},
        {"name": "Frühjahrsferien 2026", "start": "2026-02-02", "end": "2026-02-13"},
        {"name": "Osterferien 2026", "start": "2026-03-06", "end": "2026-03-20"},
        {"name": "Pfingstferien 2026", "start": "2026-05-11", "end": "2026-05-15"},
        {"name": "Sommerferien 2026", "start": "2026-07-23", "end": "2026-09-02"},
    ],
    "rheinland-pfalz": [
        {"name": "Herbstferien 2025", "start": "2025-10-13", "end": "2025-10-24"},
        {"name": "Weihnachtsferien 2025/26", "start": "2025-12-22", "end": "2026-01-06"},
        {"name": "Osterferien 2026", "start": "2026-03-23", "end": "2026-04-06"},
        {"name": "Pfingstferien 2026", "start": "2026-06-02", "end": "2026-06-10"},
        {"name": "Sommerferien 2026", "start": "2026-07-06", "end": "2026-08-14"},
    ],
}

GERMAN_PUBLIC_HOLIDAYS_2025_2026 = [
    {"name": "Neujahr", "date": "2025-01-01"},
    {"name": "Karfreitag", "date": "2025-04-18"},
    {"name": "Ostermontag", "date": "2025-04-21"},
    {"name": "Tag der Arbeit", "date": "2025-05-01"},
    {"name": "Christi Himmelfahrt", "date": "2025-05-29"},
    {"name": "Pfingstmontag", "date": "2025-06-09"},
    {"name": "Tag der Deutschen Einheit", "date": "2025-10-03"},
    {"name": "1. Weihnachtstag", "date": "2025-12-25"},
    {"name": "2. Weihnachtstag", "date": "2025-12-26"},
    {"name": "Neujahr", "date": "2026-01-01"},
    {"name": "Karfreitag", "date": "2026-04-03"},
    {"name": "Ostermontag", "date": "2026-04-06"},
    {"name": "Tag der Arbeit", "date": "2026-05-01"},
    {"name": "Christi Himmelfahrt", "date": "2026-05-14"},
    {"name": "Pfingstmontag", "date": "2026-05-25"},
    {"name": "Tag der Deutschen Einheit", "date": "2026-10-03"},
    {"name": "1. Weihnachtstag", "date": "2026-12-25"},
    {"name": "2. Weihnachtstag", "date": "2026-12-26"},
]

BUNDESLAENDER = [
    {"id": "bayern", "name": "Bayern"},
    {"id": "nrw", "name": "Nordrhein-Westfalen"},
    {"id": "berlin", "name": "Berlin"},
    {"id": "baden-wuerttemberg", "name": "Baden-Württemberg"},
    {"id": "hessen", "name": "Hessen"},
    {"id": "sachsen", "name": "Sachsen"},
    {"id": "niedersachsen", "name": "Niedersachsen"},
    {"id": "hamburg", "name": "Hamburg"},
    {"id": "rheinland-pfalz", "name": "Rheinland-Pfalz"},
]

# ============== MODELS ==============

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str
    invitation_code: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: str
    email: str
    name: str
    created_at: str
    bundesland: Optional[str] = "rheinland-pfalz"
    theme: Optional[str] = "dark"

class UserSettingsUpdate(BaseModel):
    name: Optional[str] = None
    bundesland: Optional[str] = None
    theme: Optional[str] = None

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

class SchoolYearCreate(BaseModel):
    name: str
    semester: str
    start_date: str
    end_date: str

class SchoolYearResponse(BaseModel):
    id: str
    user_id: str
    name: str
    semester: str
    start_date: str
    end_date: str
    created_at: str

class ClassSubjectCreate(BaseModel):
    name: str
    subject: str
    color: str = "#3b82f6"
    hours_per_week: int = 4
    school_year_id: str
    schedule: Optional[Dict[str, List[int]]] = None  # e.g. {"monday": [3, 4], "wednesday": [1]}

class ClassSubjectResponse(BaseModel):
    id: str
    user_id: str
    name: str
    subject: str
    color: str
    hours_per_week: int
    school_year_id: str
    created_at: str
    schedule: Optional[Dict[str, List[int]]] = None

class LessonCreate(BaseModel):
    class_subject_id: str
    date: str
    period: Optional[int] = None  # Stundennummer (1-10)
    topic: str = ""
    objective: str = ""
    curriculum_reference: str = ""
    educational_standards: str = ""
    key_terms: str = ""
    notes: str = ""
    teaching_units: int = 1
    is_cancelled: bool = False
    cancellation_reason: str = ""

class LessonResponse(BaseModel):
    id: str
    user_id: str
    class_subject_id: str
    date: str
    period: Optional[int] = None
    topic: str
    objective: str
    curriculum_reference: str
    educational_standards: str
    key_terms: str
    notes: str
    teaching_units: int
    is_cancelled: bool
    cancellation_reason: str
    created_at: str
    updated_at: str

class LessonUpdate(BaseModel):
    topic: Optional[str] = None
    objective: Optional[str] = None
    curriculum_reference: Optional[str] = None
    educational_standards: Optional[str] = None
    key_terms: Optional[str] = None
    notes: Optional[str] = None
    teaching_units: Optional[int] = None
    is_cancelled: Optional[bool] = None
    cancellation_reason: Optional[str] = None
    date: Optional[str] = None
    period: Optional[int] = None

class WorkplanEntry(BaseModel):
    date: str
    period: int
    unterrichtseinheit: Optional[str] = ""
    lehrplan: Optional[str] = ""
    stundenthema: Optional[str] = ""
    class_subject_id: Optional[str] = None

class WorkplanBulkSave(BaseModel):
    entries: List[WorkplanEntry]

class BatchLessonCreate(BaseModel):
    class_subject_id: str
    dates: List[str]
    topic: str = ""
    objective: str = ""
    curriculum_reference: str = ""
    teaching_units: int = 1

class HolidayCreate(BaseModel):
    school_year_id: str
    start_date: str
    end_date: str
    name: str

class HolidayResponse(BaseModel):
    id: str
    user_id: str
    school_year_id: str
    start_date: str
    end_date: str
    name: str

class StatisticsResponse(BaseModel):
    total_available_hours: int
    used_hours: int
    remaining_hours: int
    hours_by_weekday: Dict[str, int]
    cancelled_hours: int
    completion_percentage: float
    topics_covered: int
    hours_per_week: int
    school_weeks: int
    holiday_weeks: int
    semester_name: str
    upcoming_lessons: List[Dict[str, Any]] = []
    workplan_entries_count: int = 0

class AITopicSuggestionRequest(BaseModel):
    subject: str
    grade: str
    curriculum_topic: str
    previous_topics: List[str] = []

class AITopicSuggestionResponse(BaseModel):
    suggestions: List[Dict[str, str]]

# Sharing Models
class ShareCreate(BaseModel):
    class_subject_id: str
    shared_with_email: EmailStr
    can_edit: bool = False

class ShareResponse(BaseModel):
    id: str
    class_subject_id: str
    owner_id: str
    owner_name: str
    shared_with_id: str
    shared_with_email: str
    can_edit: bool
    created_at: str

class SharedClassResponse(BaseModel):
    id: str
    name: str
    subject: str
    color: str
    hours_per_week: int
    school_year_id: str
    owner_name: str
    owner_email: str
    can_edit: bool

# Notification Models
class NotificationResponse(BaseModel):
    id: str
    user_id: str
    type: str
    title: str
    message: str
    class_name: str
    from_user_name: str
    is_read: bool
    created_at: str

# Template Models
class TemplateCreate(BaseModel):
    name: str
    subject: str
    topic: str
    objective: str = ""
    curriculum_reference: str = ""
    educational_standards: str = ""
    key_terms: str = ""
    notes: str = ""
    teaching_units: int = 1

class TemplateResponse(BaseModel):
    id: str
    user_id: str
    name: str
    subject: str
    topic: str
    objective: str
    curriculum_reference: str
    educational_standards: str
    key_terms: str
    notes: str
    teaching_units: int
    created_at: str
    use_count: int

# Comment Models
class CommentCreate(BaseModel):
    lesson_id: str
    text: str

class CommentResponse(BaseModel):
    id: str
    lesson_id: str
    user_id: str
    user_name: str
    text: str
    created_at: str

# To-Do Models
class TodoCreate(BaseModel):
    title: str
    description: str = ""
    due_date: Optional[str] = None
    class_subject_id: Optional[str] = None
    lesson_id: Optional[str] = None
    priority: str = "medium"

class TodoResponse(BaseModel):
    id: str
    user_id: str
    title: str
    description: str
    due_date: Optional[str]
    class_subject_id: Optional[str]
    lesson_id: Optional[str]
    priority: str
    is_completed: bool
    created_at: str

class TodoUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    due_date: Optional[str] = None
    priority: Optional[str] = None
    is_completed: Optional[bool] = None

# History Models
class HistoryResponse(BaseModel):
    id: str
    user_id: str
    user_name: str
    action: str
    entity_type: str
    entity_id: str
    details: str
    created_at: str

# ============== AUTH HELPERS ==============

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed.encode())

def create_token(user_id: str, email: str) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        return user_id
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

# Helper to create notifications
async def create_notification(user_id: str, notification_type: str, title: str, message: str, class_name: str, from_user_name: str):
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "type": notification_type,
        "title": title,
        "message": message,
        "class_name": class_name,
        "from_user_name": from_user_name,
        "is_read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.notifications.insert_one(doc)
    return doc

# Helper to log history
async def log_history(user_id: str, action: str, entity_type: str, entity_id: str, details: str):
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "name": 1})
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "user_name": user.get("name", "Unbekannt") if user else "Unbekannt",
        "action": action,
        "entity_type": entity_type,
        "entity_id": entity_id,
        "details": details,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.history.insert_one(doc)

# ============== AUTH ROUTES ==============

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user: UserCreate):
    # Verify invitation code
    if user.invitation_code != INVITATION_CODE:
        raise HTTPException(status_code=403, detail="Ungültiger Einladungs-Code")
    
    existing = await db.users.find_one({"email": user.email})
    if existing:
        raise HTTPException(status_code=400, detail="E-Mail bereits registriert")
    
    user_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    user_doc = {
        "id": user_id,
        "email": user.email,
        "password": hash_password(user.password),
        "name": user.name,
        "bundesland": "rheinland-pfalz",
        "theme": "dark",
        "created_at": now
    }
    await db.users.insert_one(user_doc)
    
    token = create_token(user_id, user.email)
    return TokenResponse(
        access_token=token,
        user=UserResponse(id=user_id, email=user.email, name=user.name, created_at=now, bundesland="rheinland-pfalz", theme="dark")
    )

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_token(user["id"], user["email"])
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user["id"], 
            email=user["email"], 
            name=user["name"], 
            created_at=user["created_at"],
            bundesland=user.get("bundesland", "bayern"),
            theme=user.get("theme", "dark")
        )
    )

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(user_id: str = Depends(get_current_user)):
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return UserResponse(**user)

@api_router.put("/auth/settings", response_model=UserResponse)
async def update_settings(data: UserSettingsUpdate, user_id: str = Depends(get_current_user)):
    update_data = {k: v for k, v in data.model_dump().items() if v is not None}
    if update_data:
        await db.users.update_one({"id": user_id}, {"$set": update_data})
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    return UserResponse(**user)

# ============== GERMAN HOLIDAYS ROUTES ==============

@api_router.get("/holidays/bundeslaender")
async def get_bundeslaender():
    return BUNDESLAENDER

@api_router.get("/holidays/school-holidays/{bundesland}")
async def get_school_holidays(bundesland: str):
    if bundesland not in GERMAN_HOLIDAYS_2025_2026:
        raise HTTPException(status_code=404, detail="Bundesland nicht gefunden")
    return GERMAN_HOLIDAYS_2025_2026[bundesland]

@api_router.get("/holidays/public-holidays")
async def get_public_holidays():
    return GERMAN_PUBLIC_HOLIDAYS_2025_2026

# ============== SCHOOL YEAR ROUTES ==============

@api_router.post("/school-years", response_model=SchoolYearResponse)
async def create_school_year(data: SchoolYearCreate, user_id: str = Depends(get_current_user)):
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "name": data.name,
        "semester": data.semester,
        "start_date": data.start_date,
        "end_date": data.end_date,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.school_years.insert_one(doc)
    await log_history(user_id, "create", "school_year", doc["id"], f"Schuljahr {data.name} erstellt")
    return SchoolYearResponse(**doc)

@api_router.get("/school-years", response_model=List[SchoolYearResponse])
async def get_school_years(user_id: str = Depends(get_current_user)):
    years = await db.school_years.find({"user_id": user_id}, {"_id": 0}).to_list(100)
    return [SchoolYearResponse(**y) for y in years]

@api_router.delete("/school-years/{year_id}")
async def delete_school_year(year_id: str, user_id: str = Depends(get_current_user)):
    result = await db.school_years.delete_one({"id": year_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="School year not found")
    await db.class_subjects.delete_many({"school_year_id": year_id, "user_id": user_id})
    return {"status": "deleted"}

# ============== CLASS/SUBJECT ROUTES ==============

@api_router.post("/classes", response_model=ClassSubjectResponse)
async def create_class(data: ClassSubjectCreate, user_id: str = Depends(get_current_user)):
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "name": data.name,
        "subject": data.subject,
        "color": data.color,
        "hours_per_week": data.hours_per_week,
        "school_year_id": data.school_year_id,
        "schedule": data.schedule or {},
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.class_subjects.insert_one(doc)
    await log_history(user_id, "create", "class", doc["id"], f"Klasse {data.name} - {data.subject} erstellt")
    return ClassSubjectResponse(**doc)

@api_router.get("/classes", response_model=List[ClassSubjectResponse])
async def get_classes(school_year_id: Optional[str] = None, user_id: str = Depends(get_current_user)):
    query = {"user_id": user_id}
    if school_year_id:
        query["school_year_id"] = school_year_id
    classes = await db.class_subjects.find(query, {"_id": 0}).to_list(100)
    return [ClassSubjectResponse(**c) for c in classes]

@api_router.put("/classes/{class_id}", response_model=ClassSubjectResponse)
async def update_class(class_id: str, data: ClassSubjectCreate, user_id: str = Depends(get_current_user)):
    update_data = data.model_dump()
    result = await db.class_subjects.update_one(
        {"id": class_id, "user_id": user_id},
        {"$set": update_data}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Class not found")
    updated = await db.class_subjects.find_one({"id": class_id}, {"_id": 0})
    return ClassSubjectResponse(**updated)

@api_router.delete("/classes/{class_id}")
async def delete_class(class_id: str, user_id: str = Depends(get_current_user)):
    result = await db.class_subjects.delete_one({"id": class_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Class not found")
    await db.lessons.delete_many({"class_subject_id": class_id, "user_id": user_id})
    return {"status": "deleted"}

# ============== LESSON ROUTES ==============

@api_router.post("/lessons", response_model=LessonResponse)
async def create_lesson(data: LessonCreate, user_id: str = Depends(get_current_user)):
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "class_subject_id": data.class_subject_id,
        "date": data.date,
        "period": data.period,
        "topic": data.topic,
        "objective": data.objective,
        "curriculum_reference": data.curriculum_reference,
        "educational_standards": data.educational_standards,
        "key_terms": data.key_terms,
        "notes": data.notes,
        "teaching_units": data.teaching_units,
        "is_cancelled": data.is_cancelled,
        "cancellation_reason": data.cancellation_reason,
        "created_at": now,
        "updated_at": now
    }
    await db.lessons.insert_one(doc)
    
    # Log history
    class_info = await db.class_subjects.find_one({"id": data.class_subject_id}, {"_id": 0})
    if class_info:
        period_str = f" ({data.period}. Std.)" if data.period else ""
        await log_history(user_id, "create", "lesson", doc["id"], f"Stunde am {data.date}{period_str} für {class_info['name']} erstellt")
    
    return LessonResponse(**doc)

@api_router.post("/lessons/batch", response_model=List[LessonResponse])
async def create_batch_lessons(data: BatchLessonCreate, user_id: str = Depends(get_current_user)):
    """Create multiple lessons at once"""
    now = datetime.now(timezone.utc).isoformat()
    lessons = []
    
    for date_str in data.dates:
        doc = {
            "id": str(uuid.uuid4()),
            "user_id": user_id,
            "class_subject_id": data.class_subject_id,
            "date": date_str,
            "topic": data.topic,
            "objective": data.objective,
            "curriculum_reference": data.curriculum_reference,
            "educational_standards": "",
            "key_terms": "",
            "notes": "",
            "teaching_units": data.teaching_units,
            "is_cancelled": False,
            "cancellation_reason": "",
            "created_at": now,
            "updated_at": now
        }
        await db.lessons.insert_one(doc)
        lessons.append(LessonResponse(**doc))
    
    return lessons

# ============== WORKPLAN TABLE ROUTES ==============

@api_router.get("/workplan/{class_id}")
async def get_workplan(
    class_id: str,
    start: str = Query(...),
    end: str = Query(...),
    user_id: str = Depends(get_current_user)
):
    """Get workplan entries for a class in date range"""
    query = {
        "class_subject_id": class_id,
        "date": {"$gte": start, "$lte": end}
    }
    entries = await db.workplan.find(query, {"_id": 0}).sort([("date", 1), ("period", 1)]).to_list(1000)
    return entries

@api_router.post("/workplan/{class_id}/bulk")
async def save_workplan_bulk(
    class_id: str,
    data: WorkplanBulkSave,
    user_id: str = Depends(get_current_user)
):
    """Save multiple workplan entries at once"""
    now = datetime.now(timezone.utc).isoformat()
    
    for entry in data.entries:
        # Upsert each entry
        filter_query = {
            "class_subject_id": class_id,
            "date": entry.date,
            "period": entry.period
        }
        
        update_doc = {
            "$set": {
                "class_subject_id": class_id,
                "date": entry.date,
                "period": entry.period,
                "unterrichtseinheit": entry.unterrichtseinheit or "",
                "lehrplan": entry.lehrplan or "",
                "stundenthema": entry.stundenthema or "",
                "updated_at": now,
                "updated_by": user_id
            },
            "$setOnInsert": {
                "id": str(uuid.uuid4()),
                "created_at": now,
                "created_by": user_id
            }
        }
        
        await db.workplan.update_one(filter_query, update_doc, upsert=True)
    
    return {"status": "success", "count": len(data.entries)}

@api_router.get("/lessons", response_model=List[LessonResponse])
async def get_lessons(
    class_subject_id: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    user_id: str = Depends(get_current_user)
):
    query = {"user_id": user_id}
    if class_subject_id:
        query["class_subject_id"] = class_subject_id
    if start_date:
        query["date"] = {"$gte": start_date}
    if end_date:
        if "date" in query:
            query["date"]["$lte"] = end_date
        else:
            query["date"] = {"$lte": end_date}
    lessons = await db.lessons.find(query, {"_id": 0}).sort("date", 1).to_list(1000)
    return [LessonResponse(**l) for l in lessons]

@api_router.post("/lessons/{lesson_id}/copy", response_model=LessonResponse)
async def copy_lesson(lesson_id: str, new_date: str = Query(...), user_id: str = Depends(get_current_user)):
    """Copy an existing lesson to a new date"""
    original = await db.lessons.find_one({"id": lesson_id, "user_id": user_id}, {"_id": 0})
    if not original:
        raise HTTPException(status_code=404, detail="Stunde nicht gefunden")
    
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        **original,
        "id": str(uuid.uuid4()),
        "date": new_date,
        "created_at": now,
        "updated_at": now
    }
    await db.lessons.insert_one(doc)
    return LessonResponse(**doc)

@api_router.put("/lessons/{lesson_id}", response_model=LessonResponse)
async def update_lesson(lesson_id: str, data: LessonUpdate, user_id: str = Depends(get_current_user)):
    update_data = {k: v for k, v in data.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.lessons.update_one(
        {"id": lesson_id, "user_id": user_id},
        {"$set": update_data}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Stunde nicht gefunden")
    updated = await db.lessons.find_one({"id": lesson_id}, {"_id": 0})
    
    # Send notifications to users who have this class shared
    class_info = await db.class_subjects.find_one({"id": updated["class_subject_id"]}, {"_id": 0})
    if class_info:
        shares = await db.shares.find({"class_subject_id": updated["class_subject_id"]}, {"_id": 0}).to_list(100)
        current_user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
        class_display = f"{class_info['name']} - {class_info['subject']}"
        
        for share in shares:
            await create_notification(
                user_id=share["shared_with_id"],
                notification_type="share_edit",
                title="Arbeitsplan aktualisiert",
                message=f"{current_user['name']} hat eine Stunde im Arbeitsplan '{class_display}' geändert",
                class_name=class_display,
                from_user_name=current_user["name"]
            )
        
        await log_history(user_id, "update", "lesson", lesson_id, f"Stunde am {updated['date']} bearbeitet")
    
    return LessonResponse(**updated)

@api_router.delete("/lessons/{lesson_id}")
async def delete_lesson(lesson_id: str, user_id: str = Depends(get_current_user)):
    result = await db.lessons.delete_one({"id": lesson_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Stunde nicht gefunden")
    return {"status": "deleted"}

# ============== TEMPLATE ROUTES ==============

@api_router.post("/templates", response_model=TemplateResponse)
async def create_template(data: TemplateCreate, user_id: str = Depends(get_current_user)):
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "name": data.name,
        "subject": data.subject,
        "topic": data.topic,
        "objective": data.objective,
        "curriculum_reference": data.curriculum_reference,
        "educational_standards": data.educational_standards,
        "key_terms": data.key_terms,
        "notes": data.notes,
        "teaching_units": data.teaching_units,
        "use_count": 0,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.templates.insert_one(doc)
    return TemplateResponse(**doc)

@api_router.get("/templates", response_model=List[TemplateResponse])
async def get_templates(subject: Optional[str] = None, user_id: str = Depends(get_current_user)):
    query = {"user_id": user_id}
    if subject:
        query["subject"] = subject
    templates = await db.templates.find(query, {"_id": 0}).sort("use_count", -1).to_list(100)
    return [TemplateResponse(**t) for t in templates]

@api_router.post("/templates/{template_id}/use", response_model=TemplateResponse)
async def use_template(template_id: str, user_id: str = Depends(get_current_user)):
    """Increment use count and return template"""
    await db.templates.update_one(
        {"id": template_id, "user_id": user_id},
        {"$inc": {"use_count": 1}}
    )
    template = await db.templates.find_one({"id": template_id}, {"_id": 0})
    if not template:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")
    return TemplateResponse(**template)

@api_router.delete("/templates/{template_id}")
async def delete_template(template_id: str, user_id: str = Depends(get_current_user)):
    result = await db.templates.delete_one({"id": template_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")
    return {"status": "deleted"}

# ============== TODO ROUTES ==============

@api_router.post("/todos", response_model=TodoResponse)
async def create_todo(data: TodoCreate, user_id: str = Depends(get_current_user)):
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "title": data.title,
        "description": data.description,
        "due_date": data.due_date,
        "class_subject_id": data.class_subject_id,
        "lesson_id": data.lesson_id,
        "priority": data.priority,
        "is_completed": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.todos.insert_one(doc)
    return TodoResponse(**doc)

@api_router.get("/todos", response_model=List[TodoResponse])
async def get_todos(
    completed: Optional[bool] = None,
    class_subject_id: Optional[str] = None,
    user_id: str = Depends(get_current_user)
):
    query = {"user_id": user_id}
    if completed is not None:
        query["is_completed"] = completed
    if class_subject_id:
        query["class_subject_id"] = class_subject_id
    todos = await db.todos.find(query, {"_id": 0}).sort("due_date", 1).to_list(100)
    return [TodoResponse(**t) for t in todos]

@api_router.put("/todos/{todo_id}", response_model=TodoResponse)
async def update_todo(todo_id: str, data: TodoUpdate, user_id: str = Depends(get_current_user)):
    update_data = {k: v for k, v in data.model_dump().items() if v is not None}
    result = await db.todos.update_one(
        {"id": todo_id, "user_id": user_id},
        {"$set": update_data}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Aufgabe nicht gefunden")
    updated = await db.todos.find_one({"id": todo_id}, {"_id": 0})
    return TodoResponse(**updated)

@api_router.delete("/todos/{todo_id}")
async def delete_todo(todo_id: str, user_id: str = Depends(get_current_user)):
    result = await db.todos.delete_one({"id": todo_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Aufgabe nicht gefunden")
    return {"status": "deleted"}

# ============== COMMENT ROUTES ==============

@api_router.post("/comments", response_model=CommentResponse)
async def create_comment(data: CommentCreate, user_id: str = Depends(get_current_user)):
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "name": 1})
    doc = {
        "id": str(uuid.uuid4()),
        "lesson_id": data.lesson_id,
        "user_id": user_id,
        "user_name": user.get("name", "Unbekannt") if user else "Unbekannt",
        "text": data.text,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.comments.insert_one(doc)
    return CommentResponse(**doc)

@api_router.get("/comments/{lesson_id}", response_model=List[CommentResponse])
async def get_comments(lesson_id: str, user_id: str = Depends(get_current_user)):
    comments = await db.comments.find({"lesson_id": lesson_id}, {"_id": 0}).sort("created_at", -1).to_list(100)
    return [CommentResponse(**c) for c in comments]

@api_router.delete("/comments/{comment_id}")
async def delete_comment(comment_id: str, user_id: str = Depends(get_current_user)):
    result = await db.comments.delete_one({"id": comment_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Kommentar nicht gefunden")
    return {"status": "deleted"}

# ============== HISTORY ROUTES ==============

@api_router.get("/history", response_model=List[HistoryResponse])
async def get_history(
    entity_type: Optional[str] = None,
    entity_id: Optional[str] = None,
    limit: int = 50,
    user_id: str = Depends(get_current_user)
):
    query = {"user_id": user_id}
    if entity_type:
        query["entity_type"] = entity_type
    if entity_id:
        query["entity_id"] = entity_id
    history = await db.history.find(query, {"_id": 0}).sort("created_at", -1).to_list(limit)
    return [HistoryResponse(**h) for h in history]

@api_router.get("/history/class/{class_subject_id}", response_model=List[HistoryResponse])
async def get_class_history(class_subject_id: str, user_id: str = Depends(get_current_user)):
    """Get history for a specific class (including shared)"""
    # Get lessons for this class
    lessons = await db.lessons.find({"class_subject_id": class_subject_id}, {"id": 1, "_id": 0}).to_list(1000)
    lesson_ids = [l["id"] for l in lessons]
    
    history = await db.history.find({
        "$or": [
            {"entity_type": "class", "entity_id": class_subject_id},
            {"entity_type": "lesson", "entity_id": {"$in": lesson_ids}}
        ]
    }, {"_id": 0}).sort("created_at", -1).to_list(100)
    
    return [HistoryResponse(**h) for h in history]

# ============== SEARCH ROUTES ==============

@api_router.get("/search")
async def global_search(q: str = Query(..., min_length=2), user_id: str = Depends(get_current_user)):
    """Search across lessons, classes, and templates"""
    results = {
        "lessons": [],
        "classes": [],
        "templates": [],
        "todos": []
    }
    
    # Search lessons
    lessons = await db.lessons.find({
        "user_id": user_id,
        "$or": [
            {"topic": {"$regex": q, "$options": "i"}},
            {"key_terms": {"$regex": q, "$options": "i"}},
            {"notes": {"$regex": q, "$options": "i"}}
        ]
    }, {"_id": 0}).limit(10).to_list(10)
    results["lessons"] = lessons
    
    # Search classes
    classes = await db.class_subjects.find({
        "user_id": user_id,
        "$or": [
            {"name": {"$regex": q, "$options": "i"}},
            {"subject": {"$regex": q, "$options": "i"}}
        ]
    }, {"_id": 0}).limit(5).to_list(5)
    results["classes"] = classes
    
    # Search templates
    templates = await db.templates.find({
        "user_id": user_id,
        "$or": [
            {"name": {"$regex": q, "$options": "i"}},
            {"topic": {"$regex": q, "$options": "i"}},
            {"subject": {"$regex": q, "$options": "i"}}
        ]
    }, {"_id": 0}).limit(5).to_list(5)
    results["templates"] = templates
    
    # Search todos
    todos = await db.todos.find({
        "user_id": user_id,
        "title": {"$regex": q, "$options": "i"}
    }, {"_id": 0}).limit(5).to_list(5)
    results["todos"] = todos
    
    return results

# ============== HOLIDAY ROUTES ==============

@api_router.post("/holidays", response_model=HolidayResponse)
async def create_holiday(data: HolidayCreate, user_id: str = Depends(get_current_user)):
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "school_year_id": data.school_year_id,
        "start_date": data.start_date,
        "end_date": data.end_date,
        "name": data.name
    }
    await db.holidays.insert_one(doc)
    return HolidayResponse(**doc)

@api_router.get("/holidays", response_model=List[HolidayResponse])
async def get_holidays(school_year_id: Optional[str] = None, user_id: str = Depends(get_current_user)):
    query = {"user_id": user_id}
    if school_year_id:
        query["school_year_id"] = school_year_id
    holidays = await db.holidays.find(query, {"_id": 0}).to_list(100)
    return [HolidayResponse(**h) for h in holidays]

@api_router.delete("/holidays/{holiday_id}")
async def delete_holiday(holiday_id: str, user_id: str = Depends(get_current_user)):
    result = await db.holidays.delete_one({"id": holiday_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Holiday not found")
    return {"status": "deleted"}

# ============== STATISTICS ROUTES ==============

@api_router.get("/statistics/{class_subject_id}", response_model=StatisticsResponse)
async def get_statistics(class_subject_id: str, user_id: str = Depends(get_current_user)):
    class_info = await db.class_subjects.find_one({"id": class_subject_id, "user_id": user_id}, {"_id": 0})
    if not class_info:
        raise HTTPException(status_code=404, detail="Class not found")
    
    school_year = await db.school_years.find_one({"id": class_info["school_year_id"]}, {"_id": 0})
    if not school_year:
        raise HTTPException(status_code=404, detail="School year not found")
    
    # Calculate hours per week from schedule
    schedule = class_info.get("schedule", {})
    hours_per_week = 0
    for day, periods in schedule.items():
        if isinstance(periods, list):
            hours_per_week += len(periods)
    
    # If no schedule, fallback to hours_per_week field or default
    if hours_per_week == 0:
        hours_per_week = class_info.get("hours_per_week", 3)
    
    # Calculate school weeks
    start = datetime.fromisoformat(school_year["start_date"])
    end = datetime.fromisoformat(school_year["end_date"])
    total_weeks = (end - start).days // 7
    
    # Get holidays in this school year
    holidays = await db.holidays.find({
        "user_id": user_id,
        "school_year_id": school_year["id"]
    }, {"_id": 0}).to_list(100)
    
    # Calculate holiday weeks (approximate)
    holiday_days = 0
    for holiday in holidays:
        h_start = datetime.fromisoformat(holiday["start_date"])
        h_end = datetime.fromisoformat(holiday["end_date"])
        holiday_days += (h_end - h_start).days + 1
    holiday_weeks = holiday_days // 7
    
    # Effective school weeks
    school_weeks = total_weeks - holiday_weeks
    
    # Total available hours in semester
    total_available = school_weeks * hours_per_week
    
    # Get lessons
    lessons = await db.lessons.find({"class_subject_id": class_subject_id, "user_id": user_id}, {"_id": 0}).to_list(1000)
    
    # Get workplan entries
    workplan_entries = await db.workplan.find({"class_subject_id": class_subject_id}, {"_id": 0}).to_list(1000)
    workplan_with_content = [w for w in workplan_entries if w.get("unterrichtseinheit") or w.get("lehrplan") or w.get("stundenthema")]
    
    # Count used hours (lessons + workplan entries without duplicates)
    lesson_dates_periods = set()
    for l in lessons:
        if not l.get("is_cancelled", False):
            key = f"{l['date']}-{l.get('period', 0)}"
            lesson_dates_periods.add(key)
    
    # Hours from lessons
    used_hours = len(lesson_dates_periods)
    
    # Add workplan entries that don't have a corresponding lesson
    for w in workplan_with_content:
        key = f"{w['date']}-{w.get('period', 0)}"
        if key not in lesson_dates_periods:
            used_hours += 1
    
    cancelled_hours = len([l for l in lessons if l.get("is_cancelled", False)])
    
    # Count unique teaching units (Unterrichtseinheiten)
    unterrichtseinheiten = set()
    for w in workplan_with_content:
        if w.get("unterrichtseinheit"):
            unterrichtseinheiten.add(w["unterrichtseinheit"].strip().lower())
    for l in lessons:
        if l.get("topic") and not l.get("is_cancelled", False):
            unterrichtseinheiten.add(l["topic"].strip().lower())
    topics_covered = len(unterrichtseinheiten)
    
    # Hours by weekday
    weekday_names = ["Montag", "Dienstag", "Mittwoch", "Donnerstag", "Freitag", "Samstag", "Sonntag"]
    hours_by_weekday = {day: 0 for day in weekday_names}
    
    for lesson in lessons:
        if not lesson.get("is_cancelled", False):
            try:
                date = datetime.fromisoformat(lesson["date"])
                weekday = weekday_names[date.weekday()]
                hours_by_weekday[weekday] += 1
            except:
                pass
    
    for wp in workplan_with_content:
        try:
            date = datetime.fromisoformat(wp["date"])
            weekday = weekday_names[date.weekday()]
            hours_by_weekday[weekday] += 1
        except:
            pass
    
    # Completion percentage
    completion_percentage = (used_hours / total_available * 100) if total_available > 0 else 0
    
    # Upcoming lessons (next 5)
    today = datetime.now(timezone.utc).date().isoformat()
    upcoming = []
    
    # Combine lessons and workplan entries
    all_entries = []
    for l in lessons:
        if l["date"] >= today and not l.get("is_cancelled", False):
            all_entries.append({
                "date": l["date"],
                "period": l.get("period"),
                "topic": l.get("topic", ""),
                "source": "lesson"
            })
    for w in workplan_with_content:
        if w["date"] >= today:
            all_entries.append({
                "date": w["date"],
                "period": w.get("period"),
                "topic": w.get("unterrichtseinheit", ""),
                "source": "workplan"
            })
    
    # Sort by date and period, take first 5
    all_entries.sort(key=lambda x: (x["date"], x.get("period") or 99))
    upcoming = all_entries[:5]
    
    return StatisticsResponse(
        total_available_hours=total_available,
        used_hours=used_hours,
        remaining_hours=max(0, total_available - used_hours - cancelled_hours),
        hours_by_weekday=hours_by_weekday,
        cancelled_hours=cancelled_hours,
        completion_percentage=round(min(100, completion_percentage), 1),
        topics_covered=topics_covered,
        hours_per_week=hours_per_week,
        school_weeks=school_weeks,
        holiday_weeks=holiday_weeks,
        semester_name=school_year.get("semester", school_year.get("name", "Schuljahr")),
        upcoming_lessons=upcoming,
        workplan_entries_count=len(workplan_with_content)
    )

# ============== EXPORT ROUTES ==============

@api_router.get("/export/excel/{class_subject_id}")
async def export_excel(class_subject_id: str, user_id: str = Depends(get_current_user)):
    class_info = await db.class_subjects.find_one({"id": class_subject_id, "user_id": user_id}, {"_id": 0})
    if not class_info:
        raise HTTPException(status_code=404, detail="Class not found")
    
    lessons = await db.lessons.find({"class_subject_id": class_subject_id, "user_id": user_id}, {"_id": 0}).sort("date", 1).to_list(1000)
    
    wb = Workbook()
    ws = wb.active
    ws.title = f"{class_info['name']} - {class_info['subject']}"
    
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    
    headers = ["Datum", "Tag", "Ausfall", "Stundenthema", "Zielsetzung", "Lehrplan", "Begriffe", "UE"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
    
    weekday_names = ["Mo", "Di", "Mi", "Do", "Fr", "Sa", "So"]
    for row, lesson in enumerate(lessons, 2):
        date = datetime.fromisoformat(lesson["date"])
        ws.cell(row=row, column=1, value=date.strftime("%d.%m.%Y"))
        ws.cell(row=row, column=2, value=weekday_names[date.weekday()])
        ws.cell(row=row, column=3, value="x" if lesson["is_cancelled"] else "")
        ws.cell(row=row, column=4, value=lesson["topic"])
        ws.cell(row=row, column=5, value=lesson["objective"])
        ws.cell(row=row, column=6, value=lesson["curriculum_reference"])
        ws.cell(row=row, column=7, value=lesson["key_terms"])
        ws.cell(row=row, column=8, value=lesson["teaching_units"])
    
    for col in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    filename = f"Arbeitsplan_{class_info['name']}_{class_info['subject']}.xlsx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.get("/export/word/{class_subject_id}")
async def export_word(class_subject_id: str, user_id: str = Depends(get_current_user)):
    class_info = await db.class_subjects.find_one({"id": class_subject_id, "user_id": user_id}, {"_id": 0})
    if not class_info:
        raise HTTPException(status_code=404, detail="Class not found")
    
    school_year = await db.school_years.find_one({"id": class_info["school_year_id"]}, {"_id": 0})
    lessons = await db.lessons.find({"class_subject_id": class_subject_id, "user_id": user_id}, {"_id": 0}).sort("date", 1).to_list(1000)
    
    doc = Document()
    
    title = doc.add_heading(f"Arbeitsplan: {class_info['name']} - {class_info['subject']}", 0)
    if school_year:
        doc.add_paragraph(f"Schuljahr: {school_year['name']} ({school_year['semester']})")
    
    doc.add_paragraph("")
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    headers = ["Datum", "Thema", "Zielsetzung", "Lehrplan", "Begriffe", "UE"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    weekday_names = ["Mo", "Di", "Mi", "Do", "Fr", "Sa", "So"]
    for lesson in lessons:
        row_cells = table.add_row().cells
        date = datetime.fromisoformat(lesson["date"])
        row_cells[0].text = f"{date.strftime('%d.%m')} ({weekday_names[date.weekday()]})"
        row_cells[1].text = lesson["topic"] + (" [AUSFALL]" if lesson["is_cancelled"] else "")
        row_cells[2].text = lesson["objective"]
        row_cells[3].text = lesson["curriculum_reference"]
        row_cells[4].text = lesson["key_terms"]
        row_cells[5].text = str(lesson["teaching_units"])
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    filename = f"Arbeitsplan_{class_info['name']}_{class_info['subject']}.docx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.get("/export/pdf/{class_subject_id}")
async def export_pdf(class_subject_id: str, user_id: str = Depends(get_current_user)):
    class_info = await db.class_subjects.find_one({"id": class_subject_id, "user_id": user_id}, {"_id": 0})
    if not class_info:
        raise HTTPException(status_code=404, detail="Class not found")
    
    lessons = await db.lessons.find({"class_subject_id": class_subject_id, "user_id": user_id}, {"_id": 0}).sort("date", 1).to_list(1000)
    
    output = BytesIO()
    c = canvas.Canvas(output, pagesize=A4)
    width, height = A4
    
    c.setFont("Helvetica-Bold", 16)
    c.drawString(2*cm, height - 2*cm, f"Arbeitsplan: {class_info['name']} - {class_info['subject']}")
    
    y = height - 4*cm
    c.setFont("Helvetica-Bold", 10)
    c.drawString(2*cm, y, "Datum")
    c.drawString(4*cm, y, "Thema")
    c.drawString(12*cm, y, "UE")
    
    c.setFont("Helvetica", 9)
    y -= 0.7*cm
    
    weekday_names = ["Mo", "Di", "Mi", "Do", "Fr", "Sa", "So"]
    for lesson in lessons:
        if y < 2*cm:
            c.showPage()
            y = height - 2*cm
            c.setFont("Helvetica", 9)
        
        date = datetime.fromisoformat(lesson["date"])
        date_str = f"{date.strftime('%d.%m')} ({weekday_names[date.weekday()]})"
        topic = lesson["topic"][:50] + "..." if len(lesson["topic"]) > 50 else lesson["topic"]
        if lesson["is_cancelled"]:
            topic = "[AUSFALL] " + topic
        
        c.drawString(2*cm, y, date_str)
        c.drawString(4*cm, y, topic)
        c.drawString(12*cm, y, str(lesson["teaching_units"]))
        y -= 0.5*cm
    
    c.save()
    output.seek(0)
    
    filename = f"Arbeitsplan_{class_info['name']}_{class_info['subject']}.pdf"
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ============== AI SUGGESTIONS ROUTE ==============

@api_router.post("/ai/suggestions", response_model=AITopicSuggestionResponse)
async def get_ai_suggestions(data: AITopicSuggestionRequest, user_id: str = Depends(get_current_user)):
    import asyncio
    import json
    import re
    
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        api_key = os.environ.get('EMERGENT_LLM_KEY')
        if not api_key:
            raise HTTPException(status_code=500, detail="AI service not configured")
        
        chat = LlmChat(
            api_key=api_key,
            session_id=f"planed-{user_id}-{uuid.uuid4()}",
            system_message="""Du bist ein erfahrener Lehrer und Lehrplanexperte. 
            Erstelle konkrete, praxisnahe Unterrichtsvorschläge basierend auf dem deutschen Lehrplan.
            Antworte immer auf Deutsch und strukturiere deine Vorschläge klar."""
        ).with_model("gemini", "gemini-3-flash-preview")
        
        prompt = f"""Erstelle 3 Unterrichtsthemen-Vorschläge für:
        Fach: {data.subject}
        Klassenstufe: {data.grade}
        Lehrplanthema: {data.curriculum_topic}
        {"Bereits behandelte Themen: " + ", ".join(data.previous_topics[:3]) if data.previous_topics else ""}
        
        Antworte NUR mit einem JSON-Array: [{{"topic": "...", "objective": "...", "key_terms": "..."}}]"""
        
        try:
            response = await asyncio.wait_for(
                chat.send_message(UserMessage(text=prompt)),
                timeout=25.0
            )
        except asyncio.TimeoutError:
            return AITopicSuggestionResponse(suggestions=[
                {"topic": f"Einführung in {data.curriculum_topic}", "objective": "Grundlagen verstehen", "key_terms": data.subject},
                {"topic": f"Vertiefung: {data.curriculum_topic}", "objective": "Anwendung üben", "key_terms": "Übung, Praxis"},
                {"topic": f"Zusammenfassung {data.curriculum_topic}", "objective": "Wissen festigen", "key_terms": "Wiederholung"}
            ])
        
        json_match = re.search(r'\[.*\]', response, re.DOTALL)
        if json_match:
            try:
                suggestions = json.loads(json_match.group())
            except json.JSONDecodeError:
                suggestions = [{"topic": "Vorschlag konnte nicht generiert werden", "objective": "", "key_terms": ""}]
        else:
            suggestions = [{"topic": "Vorschlag konnte nicht generiert werden", "objective": "", "key_terms": ""}]
        
        return AITopicSuggestionResponse(suggestions=suggestions[:5])
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI suggestion error: {e}")
        return AITopicSuggestionResponse(suggestions=[
            {"topic": f"Thema zu {data.curriculum_topic}", "objective": "Lernziele definieren", "key_terms": data.subject}
        ])

# ============== DOCUMENT MANAGEMENT ==============

@api_router.post("/documents/upload")
async def upload_document(
    class_subject_id: str,
    lesson_id: Optional[str] = None,
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user)
):
    allowed_types = [".docx", ".doc", ".pdf", ".jpg", ".jpeg", ".png"]
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in allowed_types:
        raise HTTPException(status_code=400, detail=f"File type not allowed. Allowed: {allowed_types}")
    
    content = await file.read()
    
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "class_subject_id": class_subject_id,
        "lesson_id": lesson_id,
        "filename": file.filename,
        "content_type": file.content_type,
        "size": len(content),
        "content": content,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.documents.insert_one(doc)
    
    return {"id": doc["id"], "filename": file.filename, "size": len(content)}

@api_router.get("/documents")
async def get_documents(
    class_subject_id: Optional[str] = None,
    lesson_id: Optional[str] = None,
    user_id: str = Depends(get_current_user)
):
    query = {"user_id": user_id}
    if class_subject_id:
        query["class_subject_id"] = class_subject_id
    if lesson_id:
        query["lesson_id"] = lesson_id
    
    docs = await db.documents.find(query, {"_id": 0, "content": 0}).to_list(100)
    return docs

@api_router.get("/documents/{doc_id}/download")
async def download_document(doc_id: str, user_id: str = Depends(get_current_user)):
    doc = await db.documents.find_one({"id": doc_id, "user_id": user_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return StreamingResponse(
        BytesIO(doc["content"]),
        media_type=doc["content_type"],
        headers={"Content-Disposition": f"attachment; filename={doc['filename']}"}
    )

@api_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, user_id: str = Depends(get_current_user)):
    result = await db.documents.delete_one({"id": doc_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "deleted"}

# ============== SHARING ROUTES ==============

@api_router.post("/shares", response_model=ShareResponse)
async def share_class(data: ShareCreate, user_id: str = Depends(get_current_user)):
    class_info = await db.class_subjects.find_one({"id": data.class_subject_id, "user_id": user_id}, {"_id": 0})
    if not class_info:
        raise HTTPException(status_code=404, detail="Klasse nicht gefunden")
    
    target_user = await db.users.find_one({"email": data.shared_with_email}, {"_id": 0, "password": 0})
    if not target_user:
        raise HTTPException(status_code=404, detail="Benutzer mit dieser E-Mail nicht gefunden")
    
    if target_user["id"] == user_id:
        raise HTTPException(status_code=400, detail="Sie können nicht mit sich selbst teilen")
    
    existing = await db.shares.find_one({
        "class_subject_id": data.class_subject_id,
        "shared_with_id": target_user["id"]
    })
    if existing:
        raise HTTPException(status_code=400, detail="Bereits mit diesem Benutzer geteilt")
    
    owner = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    
    share_doc = {
        "id": str(uuid.uuid4()),
        "class_subject_id": data.class_subject_id,
        "owner_id": user_id,
        "owner_name": owner["name"],
        "shared_with_id": target_user["id"],
        "shared_with_email": data.shared_with_email,
        "can_edit": data.can_edit,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.shares.insert_one(share_doc)
    
    class_display = f"{class_info['name']} - {class_info['subject']}"
    permission = "Bearbeitung" if data.can_edit else "Nur Ansicht"
    await create_notification(
        user_id=target_user["id"],
        notification_type="share_new",
        title="Neuer Arbeitsplan geteilt",
        message=f"{owner['name']} hat den Arbeitsplan '{class_display}' mit Ihnen geteilt ({permission})",
        class_name=class_display,
        from_user_name=owner["name"]
    )
    
    return ShareResponse(**share_doc)

@api_router.get("/shares/my-shares", response_model=List[ShareResponse])
async def get_my_shares(user_id: str = Depends(get_current_user)):
    shares = await db.shares.find({"owner_id": user_id}, {"_id": 0}).to_list(100)
    return [ShareResponse(**s) for s in shares]

@api_router.get("/shares/shared-with-me", response_model=List[SharedClassResponse])
async def get_shared_with_me(user_id: str = Depends(get_current_user)):
    shares = await db.shares.find({"shared_with_id": user_id}, {"_id": 0}).to_list(100)
    
    shared_classes = []
    for share in shares:
        class_info = await db.class_subjects.find_one({"id": share["class_subject_id"]}, {"_id": 0})
        if class_info:
            owner = await db.users.find_one({"id": share["owner_id"]}, {"_id": 0, "password": 0})
            shared_classes.append(SharedClassResponse(
                id=class_info["id"],
                name=class_info["name"],
                subject=class_info["subject"],
                color=class_info["color"],
                hours_per_week=class_info["hours_per_week"],
                school_year_id=class_info["school_year_id"],
                owner_name=owner["name"] if owner else "Unbekannt",
                owner_email=owner["email"] if owner else "",
                can_edit=share["can_edit"]
            ))
    
    return shared_classes

@api_router.delete("/shares/{share_id}")
async def remove_share(share_id: str, user_id: str = Depends(get_current_user)):
    result = await db.shares.delete_one({"id": share_id, "owner_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Freigabe nicht gefunden")
    return {"status": "deleted"}

@api_router.get("/shares/class/{class_subject_id}", response_model=List[ShareResponse])
async def get_class_shares(class_subject_id: str, user_id: str = Depends(get_current_user)):
    class_info = await db.class_subjects.find_one({"id": class_subject_id, "user_id": user_id}, {"_id": 0})
    if not class_info:
        raise HTTPException(status_code=404, detail="Klasse nicht gefunden")
    
    shares = await db.shares.find({"class_subject_id": class_subject_id}, {"_id": 0}).to_list(100)
    return [ShareResponse(**s) for s in shares]

# ============== NOTIFICATION ROUTES ==============

@api_router.get("/notifications", response_model=List[NotificationResponse])
async def get_notifications(user_id: str = Depends(get_current_user)):
    notifications = await db.notifications.find(
        {"user_id": user_id}, 
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    return [NotificationResponse(**n) for n in notifications]

@api_router.get("/notifications/unread-count")
async def get_unread_count(user_id: str = Depends(get_current_user)):
    count = await db.notifications.count_documents({"user_id": user_id, "is_read": False})
    return {"count": count}

@api_router.put("/notifications/{notification_id}/read")
async def mark_as_read(notification_id: str, user_id: str = Depends(get_current_user)):
    result = await db.notifications.update_one(
        {"id": notification_id, "user_id": user_id},
        {"$set": {"is_read": True}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Benachrichtigung nicht gefunden")
    return {"status": "read"}

@api_router.put("/notifications/read-all")
async def mark_all_as_read(user_id: str = Depends(get_current_user)):
    await db.notifications.update_many(
        {"user_id": user_id, "is_read": False},
        {"$set": {"is_read": True}}
    )
    return {"status": "all_read"}

@api_router.delete("/notifications/{notification_id}")
async def delete_notification(notification_id: str, user_id: str = Depends(get_current_user)):
    result = await db.notifications.delete_one({"id": notification_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Benachrichtigung nicht gefunden")
    return {"status": "deleted"}

# ============== RESEARCH API ==============

@api_router.get("/research/images")
async def search_images(query: str, user_id: str = Depends(get_current_user)):
    """Search for educational images using Wikimedia Commons API (free, no API key required)"""
    try:
        results = []
        
        # Wikimedia requires a User-Agent header
        headers = {
            "User-Agent": "PlanEd/2.0 (Educational Teacher Planning Tool; contact@planed.app)"
        }
        
        async with httpx.AsyncClient(headers=headers) as http_client:
            # Use Wikimedia Commons API - completely free, no API key needed
            response = await http_client.get(
                "https://commons.wikimedia.org/w/api.php",
                params={
                    "action": "query",
                    "format": "json",
                    "generator": "search",
                    "gsrsearch": f"filetype:bitmap {query}",
                    "gsrlimit": 15,
                    "gsrnamespace": 6,
                    "prop": "imageinfo",
                    "iiprop": "url|extmetadata|size",
                    "iiurlwidth": 400
                },
                timeout=15.0
            )
            
            if response.status_code == 200:
                data = response.json()
                pages = data.get("query", {}).get("pages", {})
                
                for page_id, page in pages.items():
                    if "imageinfo" in page and page.get("imageinfo"):
                        info = page["imageinfo"][0]
                        extmeta = info.get("extmetadata", {})
                        
                        # Get description
                        desc = extmeta.get("ImageDescription", {}).get("value", "")
                        if desc:
                            # Strip HTML tags
                            import re
                            desc = re.sub('<[^<]+?>', '', desc)[:100]
                        else:
                            desc = page.get("title", "").replace("File:", "").replace("_", " ")
                        
                        # Get author
                        author = extmeta.get("Artist", {}).get("value", "")
                        if author:
                            author = re.sub('<[^<]+?>', '', author)[:50]
                        else:
                            author = "Wikimedia Commons"
                        
                        results.append({
                            "id": str(page_id),
                            "url": info.get("descriptionurl", info.get("url", "")),
                            "thumb": info.get("thumburl", info.get("url", "")),
                            "description": desc,
                            "author": author,
                            "download_url": info.get("url", ""),
                            "source": "Wikimedia Commons"
                        })
            
            # Fallback if no results: provide direct search links
            if not results:
                from urllib.parse import quote
                url_encoded_query = quote(query)
                hyphen_query = query.replace(" ", "-").lower()
                plus_query = query.replace(" ", "+")
                results = [
                    {
                        "id": "search-wikimedia",
                        "url": f"https://commons.wikimedia.org/w/index.php?search={url_encoded_query}&title=Special:MediaSearch&type=image",
                        "thumb": None,
                        "description": f"Auf Wikimedia Commons nach '{query}' suchen",
                        "author": "Wikimedia Commons",
                        "download_url": f"https://commons.wikimedia.org/w/index.php?search={url_encoded_query}&title=Special:MediaSearch&type=image",
                        "source": "Wikimedia Commons",
                        "is_link": True
                    },
                    {
                        "id": "search-pixabay",
                        "url": f"https://pixabay.com/images/search/{url_encoded_query}/",
                        "thumb": None,
                        "description": f"Auf Pixabay nach '{query}' suchen",
                        "author": "Pixabay",
                        "download_url": f"https://pixabay.com/images/search/{url_encoded_query}/",
                        "source": "Pixabay",
                        "is_link": True
                    },
                    {
                        "id": "search-unsplash",
                        "url": f"https://unsplash.com/s/photos/{hyphen_query}",
                        "thumb": None,
                        "description": f"Auf Unsplash nach '{query}' suchen",
                        "author": "Unsplash",
                        "download_url": f"https://unsplash.com/s/photos/{hyphen_query}",
                        "source": "Unsplash",
                        "is_link": True
                    }
                ]
        
        return {"results": results, "total": len(results)}
    except Exception as e:
        logger.error(f"Image search error: {e}")
        return {"results": [], "total": 0, "error": str(e)}

@api_router.get("/research/videos")
async def search_videos(query: str, user_id: str = Depends(get_current_user)):
    """Search for educational YouTube videos using web scraping fallback"""
    try:
        results = []
        
        # Try YouTube Data API first
        api_key = os.environ.get("YOUTUBE_API_KEY", "")
        if api_key:
            async with httpx.AsyncClient() as client:
                response = await client.get(
                    "https://www.googleapis.com/youtube/v3/search",
                    params={
                        "part": "snippet",
                        "q": f"{query} Unterricht Schule",
                        "type": "video",
                        "maxResults": 10,
                        "relevanceLanguage": "de",
                        "safeSearch": "strict",
                        "key": api_key
                    },
                    timeout=10.0
                )
                
                if response.status_code == 200:
                    data = response.json()
                    for item in data.get("items", []):
                        results.append({
                            "id": item["id"]["videoId"],
                            "title": item["snippet"]["title"],
                            "description": item["snippet"]["description"][:200],
                            "thumbnail": item["snippet"]["thumbnails"]["medium"]["url"],
                            "channel": item["snippet"]["channelTitle"],
                            "url": f"https://www.youtube.com/watch?v={item['id']['videoId']}",
                            "source": "YouTube"
                        })
        
        # If no API key or no results, provide helpful message
        if not results:
            # Return curated educational channel suggestions
            educational_channels = [
                {"name": "simpleclub", "topic": "Mathematik, Physik, Chemie"},
                {"name": "MrWissen2go", "topic": "Geschichte, Politik"},
                {"name": "Duden Learnattack", "topic": "Alle Fächer"},
                {"name": "TheSimpleClub", "topic": "MINT-Fächer"},
                {"name": "musstewissen", "topic": "Deutsch, Mathe, Chemie"},
            ]
            
            return {
                "results": [],
                "total": 0,
                "message": "YouTube-Suche benötigt API-Key. Empfohlene Bildungskanäle:",
                "suggestions": educational_channels,
                "search_url": f"https://www.youtube.com/results?search_query={query.replace(' ', '+')}+Unterricht"
            }
        
        return {"results": results, "total": len(results)}
    except Exception as e:
        logger.error(f"Video search error: {e}")
        return {"results": [], "total": 0, "error": str(e)}

@api_router.get("/research/papers")
async def search_academic_papers(query: str, source: str = "semantic_scholar", user_id: str = Depends(get_current_user)):
    """Search for academic papers from Semantic Scholar or OpenAlex"""
    try:
        async with httpx.AsyncClient() as client:
            results = []
            
            if source == "semantic_scholar":
                response = await client.get(
                    "https://api.semanticscholar.org/graph/v1/paper/search",
                    params={
                        "query": query,
                        "limit": 10,
                        "fields": "title,abstract,authors,year,url,citationCount,openAccessPdf"
                    },
                    timeout=15.0
                )
                
                if response.status_code == 200:
                    data = response.json()
                    for paper in data.get("data", []):
                        results.append({
                            "id": paper.get("paperId", ""),
                            "title": paper.get("title", ""),
                            "abstract": paper.get("abstract", "")[:500] if paper.get("abstract") else "",
                            "authors": ", ".join([a.get("name", "") for a in paper.get("authors", [])[:3]]),
                            "year": paper.get("year"),
                            "citations": paper.get("citationCount", 0),
                            "url": paper.get("url", ""),
                            "pdf_url": paper.get("openAccessPdf", {}).get("url") if paper.get("openAccessPdf") else None,
                            "source": "Semantic Scholar"
                        })
            
            elif source == "openalex":
                response = await client.get(
                    "https://api.openalex.org/works",
                    params={
                        "search": query,
                        "per_page": 10,
                        "sort": "cited_by_count:desc"
                    },
                    headers={"User-Agent": "PlanEd-App/1.0"},
                    timeout=15.0
                )
                
                if response.status_code == 200:
                    data = response.json()
                    for work in data.get("results", []):
                        abstract = ""
                        if work.get("abstract_inverted_index"):
                            # Reconstruct abstract from inverted index
                            inv_idx = work["abstract_inverted_index"]
                            words = [(word, min(positions)) for word, positions in inv_idx.items()]
                            words.sort(key=lambda x: x[1])
                            abstract = " ".join([w[0] for w in words])[:500]
                        
                        results.append({
                            "id": work.get("id", "").split("/")[-1],
                            "title": work.get("title", ""),
                            "abstract": abstract,
                            "authors": ", ".join([a.get("author", {}).get("display_name", "") for a in work.get("authorships", [])[:3]]),
                            "year": work.get("publication_year"),
                            "citations": work.get("cited_by_count", 0),
                            "url": work.get("doi", "") if work.get("doi") else work.get("id", ""),
                            "pdf_url": work.get("open_access", {}).get("oa_url"),
                            "source": "OpenAlex"
                        })
            
            return {"results": results, "total": len(results)}
    except Exception as e:
        logger.error(f"Academic search error: {e}")
        return {"results": [], "total": 0, "error": str(e)}

@api_router.post("/research/translate")
async def translate_text(text: str = "", target_lang: str = "de", user_id: str = Depends(get_current_user)):
    """Translate text to German using AI"""
    if not text:
        return {"translated": "", "error": "No text provided"}
    
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        # Use Emergent LLM key for translation
        emergent_key = os.environ.get("EMERGENT_MODEL_API_KEY", "") or os.environ.get("EMERGENT_LLM_KEY", "")
        if not emergent_key:
            return {"translated": text, "error": "Translation service not configured"}
        
        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"translate-{user_id}",
            system_message="Du bist ein professioneller Übersetzer für wissenschaftliche Texte. Übersetze Texte präzise ins Deutsche. Behalte Fachbegriffe bei, wenn sie im Deutschen üblich sind. Antworte NUR mit der Übersetzung, ohne Erklärungen oder zusätzliche Kommentare."
        ).with_model("gemini", "gemini-2.0-flash")
        
        prompt = f"Übersetze folgenden wissenschaftlichen Abstract ins Deutsche:\n\n{text[:2000]}"
        
        response = await asyncio.wait_for(
            chat.send_message(UserMessage(text=prompt)),
            timeout=30.0
        )
        
        return {"translated": response, "original": text[:500]}
    except asyncio.TimeoutError:
        return {"translated": text, "error": "Translation timeout"}
    except Exception as e:
        logger.error(f"Translation error: {e}")
        return {"translated": text, "error": str(e)}

# ============== LEHRPLAN-BASIERTE UNTERRICHTSPLANUNG ==============

# Import der Fach-Daten aus Modulen
from data.lehrplan_deutsch_rlp import LEHRPLAN_DEUTSCH_RLP
from data.schulbuecher_deutsch import SCHULBUECHER_DEUTSCH

# Pydantic Models für Unterrichtsplanung (bleiben hier wegen Abhängigkeiten)
# Pydantic Models für Unterrichtsplanung
class UnterrichtsreiheRequest(BaseModel):
    klassenstufe: str
    kompetenzbereich: str
    thema_id: str
    niveau: str  # G, M, E
    stunden_anzahl: int = 6
    schulbuch_id: Optional[str] = None  # Optional: Schulbuch-Referenz

class MaterialRequest(BaseModel):
    thema: str
    niveau: str
    material_typ: str  # arbeitsblatt, quiz, raetsel, zuordnung
    klassenstufe: str

class GeneratedMaterial(BaseModel):
    typ: str
    titel: str
    inhalt: str
    loesung: Optional[str] = None

# API Endpoints


# ============== IMPORT AUSGELAGERTER MODULE ==============
from routes.faecher.deutsch import router as deutsch_router

# Include routers
app.include_router(deutsch_router)

# ============== ROOT ==============

@api_router.get("/")
async def root():
    return {"message": "PlanEd API v2.1.0 - Modular", "status": "running"}

# Include the main router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
