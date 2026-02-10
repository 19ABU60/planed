# Deutsch Fach-Modul für PlanEd
# Enthält alle Endpunkte für die Deutsch-Unterrichtsplanung

from fastapi import APIRouter, HTTPException, Depends, Query
from typing import Optional
from datetime import datetime, timezone
import asyncio
import uuid
import os
import logging

from services.auth import get_db, get_current_user
from data.lehrplan_deutsch_rlp import LEHRPLAN_DEUTSCH_RLP
from data.schulbuecher_deutsch import SCHULBUECHER_DEUTSCH

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/lehrplan", tags=["deutsch"])


# ============== Pydantic Models ==============

from pydantic import BaseModel
from typing import Dict, Any

class UnterrichtsreiheRequest(BaseModel):
    klassenstufe: str
    kompetenzbereich: str
    thema_id: str
    niveau: str  # G, M, E
    stunden_anzahl: int = 6
    schulbuch_id: Optional[str] = None

class MaterialRequest(BaseModel):
    thema: str
    niveau: str
    material_typ: str  # arbeitsblatt, quiz, raetsel, zuordnung
    klassenstufe: str


# ============== LEHRPLAN STRUKTUR ==============

@router.get("/struktur")
async def get_lehrplan_struktur(user_id: str = Depends(get_current_user)):
    """Gibt die komplette LP-Struktur für das Auswahlmenü zurück"""
    struktur = {}
    for klassenstufe, bereiche in LEHRPLAN_DEUTSCH_RLP.items():
        struktur[klassenstufe] = {}
        for bereich_id, bereich_data in bereiche.items():
            struktur[klassenstufe][bereich_id] = {
                "name": bereich_data["name"],
                "themen": [{"id": t["id"], "name": t["name"]} for t in bereich_data["themen"]]
            }
    return {"fach": "Deutsch", "bundesland": "RLP", "schulart": "RS+", "struktur": struktur}


@router.get("/thema")
async def get_thema_details(
    klassenstufe: str = Query(...),
    kompetenzbereich: str = Query(...),
    thema_id: str = Query(...),
    user_id: str = Depends(get_current_user)
):
    """Gibt Details zu einem spezifischen Thema zurück"""
    try:
        bereich = LEHRPLAN_DEUTSCH_RLP.get(klassenstufe, {}).get(kompetenzbereich, {})
        themen = bereich.get("themen", [])
        for thema in themen:
            if thema["id"] == thema_id:
                return {
                    "klassenstufe": klassenstufe,
                    "kompetenzbereich": bereich.get("name", kompetenzbereich),
                    "thema": thema
                }
        raise HTTPException(status_code=404, detail="Thema nicht gefunden")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============== SCHULBÜCHER ==============

@router.get("/schulbuecher")
async def get_schulbuecher(
    klassenstufe: Optional[str] = Query(None),
    user_id: str = Depends(get_current_user)
):
    """Gibt verfügbare Schulbücher zurück, optional gefiltert nach Klassenstufe"""
    try:
        result = []
        for buch_id, buch in SCHULBUECHER_DEUTSCH.items():
            if klassenstufe and buch["klassenstufe"] != "alle":
                if klassenstufe not in buch["klassenstufe"]:
                    continue
            result.append({
                "id": buch["id"],
                "name": buch["name"],
                "verlag": buch["verlag"],
                "isbn": buch["isbn"],
                "klassenstufe": buch["klassenstufe"],
                "kapitel": list(buch["kapitel"].keys())
            })
        return {"schulbuecher": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============== UNTERRICHTSREIHE GENERIEREN ==============

@router.post("/unterrichtsreihe/generieren")
async def generiere_unterrichtsreihe(
    request: UnterrichtsreiheRequest,
    user_id: str = Depends(get_current_user)
):
    """Generiert eine Unterrichtsreihe mit KI, optional mit Schulbuch-Referenzen"""
    import json as json_lib
    db = get_db()
    
    try:
        # Hole Thema-Details
        bereich = LEHRPLAN_DEUTSCH_RLP.get(request.klassenstufe, {}).get(request.kompetenzbereich, {})
        thema_data = None
        for t in bereich.get("themen", []):
            if t["id"] == request.thema_id:
                thema_data = t
                break
        
        if not thema_data:
            raise HTTPException(status_code=404, detail="Thema nicht gefunden")
        
        niveau_text = thema_data.get(request.niveau, "")
        niveau_name = {"G": "grundlegend", "M": "mittel", "E": "erweitert"}.get(request.niveau, "mittel")
        
        # Schulbuch-Informationen vorbereiten
        schulbuch_info = ""
        schulbuch_name = None
        if request.schulbuch_id and request.schulbuch_id != "kein_schulbuch":
            schulbuch = SCHULBUECHER_DEUTSCH.get(request.schulbuch_id)
            if schulbuch:
                schulbuch_name = f"{schulbuch['name']} ({schulbuch['verlag']})"
                kapitel_info = []
                for kap_id, kap_data in schulbuch.get("kapitel", {}).items():
                    kap_themen_lower = [t.lower() for t in kap_data.get("themen", [])]
                    bereich_name_lower = bereich.get("name", "").lower()
                    if any(word in bereich_name_lower for word in ["schreib", "erzähl", "bericht"]) and kap_id in ["erzaehlen", "berichten"]:
                        kapitel_info.append(f"- {kap_data['name']}: Seiten {kap_data['seiten']}")
                    elif any(word in bereich_name_lower for word in ["lesen", "text", "literatur"]) and kap_id == "lesen":
                        kapitel_info.append(f"- {kap_data['name']}: Seiten {kap_data['seiten']}")
                    elif any(word in bereich_name_lower for word in ["sprach", "grammatik", "rechtschreib"]) and kap_id in ["grammatik", "rechtschreibung"]:
                        kapitel_info.append(f"- {kap_data['name']}: Seiten {kap_data['seiten']}")
                    elif any(word in bereich_name_lower for word in ["argument", "diskut"]) and kap_id == "argumentieren":
                        kapitel_info.append(f"- {kap_data['name']}: Seiten {kap_data['seiten']}")
                
                if not kapitel_info:
                    for kap_id, kap_data in schulbuch.get("kapitel", {}).items():
                        kapitel_info.append(f"- {kap_data['name']}: Seiten {kap_data['seiten']}")
                
                schulbuch_info = f"""
SCHULBUCH-BEZUG:
Verwende das Schulbuch "{schulbuch['name']}" ({schulbuch['verlag']}, ISBN: {schulbuch['isbn']}).
Relevante Kapitel:
{chr(10).join(kapitel_info)}

Bei jeder Stunde: Gib konkrete Seitenzahlen an, z.B. "S. 34-36" oder "Aufgabe 3 auf S. 42".
Füge bei jedem Material-Eintrag einen Schulbuch-Verweis hinzu wenn passend."""
        
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        emergent_key = os.environ.get("EMERGENT_LLM_KEY", "")
        if not emergent_key:
            raise HTTPException(status_code=500, detail="KI-Service nicht konfiguriert")
        
        system_msg = """Du bist ein erfahrener Deutschlehrer an einer Realschule plus in Rheinland-Pfalz. 
Du erstellst praxisnahe, differenzierte Unterrichtsreihen für den Deutschunterricht.
Deine Unterrichtsreihen sind klar strukturiert, schülerorientiert und enthalten konkrete Aktivitäten.
Wenn ein Schulbuch angegeben ist, integrierst du passende Seiten und Aufgaben aus diesem Buch.
Antworte IMMER im JSON-Format."""
        
        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"unterricht-{user_id}-{uuid.uuid4()}",
            system_message=system_msg
        ).with_model("gemini", "gemini-3-flash-preview")
        
        json_format = """{
    "titel": "Titel der Unterrichtsreihe",
    "ueberblick": "Kurze Beschreibung (2-3 Sätze)",
    "schulbuch": "Name des Schulbuchs (falls verwendet)",
    "lernziele": ["Übergeordnetes Lernziel 1", "Übergeordnetes Lernziel 2"],
    "stunden": [
        {
            "nummer": 1,
            "titel": "Titel der Stunde",
            "lernziel": "Konkretes Lernziel dieser Stunde (Die SuS können...)",
            "phase": "Einstieg/Erarbeitung/Sicherung",
            "dauer": "45 min",
            "inhalt": "Detaillierte Beschreibung der Stundeninhalte",
            "aufgaben": ["Konkrete Aufgabe 1", "Konkrete Aufgabe 2"],
            "methoden": ["Methode 1", "Methode 2"],
            "material": ["Benötigtes Material"],
            "schulbuch_seiten": "z.B. S. 34-36, Aufgabe 1-3"
        }
    ],
    "differenzierung": {
        "foerdern": "Maßnahmen für schwächere SuS",
        "fordern": "Maßnahmen für stärkere SuS"
    },
    "leistungsnachweis": "Vorschlag für Leistungsüberprüfung"
}"""
        
        prompt = f"""Erstelle eine Unterrichtsreihe für Deutsch RS+ mit folgenden Parametern:

Klassenstufe: {request.klassenstufe}
Kompetenzbereich: {bereich.get('name', request.kompetenzbereich)}
Thema: {thema_data['name']}
Lehrplaninhalt ({niveau_name}): {niveau_text}
Niveau: {niveau_name} ({"einfacher, mehr Unterstützung" if request.niveau == "G" else "anspruchsvoller, mehr Eigenständigkeit" if request.niveau == "E" else "mittleres Anforderungsniveau"})
Anzahl Unterrichtsstunden: {request.stunden_anzahl}
{schulbuch_info}

Erstelle eine detaillierte Unterrichtsreihe im folgenden JSON-Format:
{json_format}

Wichtig: 
- Genau {request.stunden_anzahl} Stunden erstellen
- Niveau {niveau_name} beachten
- JEDE Stunde muss ein konkretes "lernziel" haben (beginnt mit "Die SuS können...")
- JEDE Stunde muss 2-4 konkrete "aufgaben" haben (z.B. "Textvergleich", "Tabelle erstellen", "Partnerarbeit")
- Praxisnah und umsetzbar
{"- Bei JEDER Stunde konkrete Schulbuch-Seitenzahlen angeben!" if schulbuch_info else "- schulbuch_seiten kann leer bleiben wenn kein Schulbuch gewählt"}
- Nur valides JSON zurückgeben, keine Erklärungen davor oder danach"""

        response = await asyncio.wait_for(
            chat.send_message(UserMessage(text=prompt)),
            timeout=60.0
        )
        
        response_text = response.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        unterrichtsreihe = json_lib.loads(response_text.strip())
        
        doc = {
            "user_id": user_id,
            "klassenstufe": request.klassenstufe,
            "kompetenzbereich": request.kompetenzbereich,
            "thema_id": request.thema_id,
            "niveau": request.niveau,
            "schulbuch_id": request.schulbuch_id,
            "schulbuch_name": schulbuch_name,
            "unterrichtsreihe": unterrichtsreihe,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        result = await db.unterrichtsreihen.insert_one(doc)
        
        return {
            "id": str(result.inserted_id),
            "unterrichtsreihe": unterrichtsreihe,
            "schulbuch": schulbuch_name,
            "meta": {
                "klassenstufe": request.klassenstufe,
                "thema": thema_data["name"],
                "niveau": niveau_name
            }
        }
        
    except json_lib.JSONDecodeError as e:
        logger.error(f"JSON Parse error: {e}")
        raise HTTPException(status_code=500, detail="Fehler beim Parsen der KI-Antwort")
    except asyncio.TimeoutError:
        raise HTTPException(status_code=504, detail="KI-Anfrage Timeout")
    except Exception as e:
        logger.error(f"Unterrichtsreihe generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============== MATERIAL GENERIEREN ==============

@router.post("/material/generieren")
async def generiere_material(
    request: MaterialRequest,
    user_id: str = Depends(get_current_user)
):
    """Generiert Unterrichtsmaterial (Arbeitsblatt, Quiz, Rätsel) mit KI"""
    import json as json_lib
    
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        emergent_key = os.environ.get("EMERGENT_LLM_KEY", "")
        if not emergent_key:
            raise HTTPException(status_code=500, detail="KI-Service nicht konfiguriert")
        
        niveau_name = {"G": "grundlegend", "M": "mittel", "E": "erweitert"}.get(request.niveau, "mittel")
        
        material_prompts = {
            "arbeitsblatt": f"""Erstelle ein Arbeitsblatt für Deutsch RS+ Klasse {request.klassenstufe} zum Thema "{request.thema}" (Niveau: {niveau_name}).

Das Arbeitsblatt soll enthalten:
- Überschrift
- Kurze Einleitung/Erklärung
- 4-6 abwechslungsreiche Aufgaben
- Platz für Schülerantworten (markiert mit ___)

Format als JSON:
{{"titel": "...", "einleitung": "...", "aufgaben": [{{"nummer": 1, "aufgabenstellung": "...", "punkte": 2}}], "loesung": [{{"nummer": 1, "loesung": "..."}}]}}""",

            "quiz": f"""Erstelle ein Quiz für Deutsch RS+ Klasse {request.klassenstufe} zum Thema "{request.thema}" (Niveau: {niveau_name}).

Das Quiz soll 8 Multiple-Choice-Fragen enthalten.

Format als JSON:
{{"titel": "Quiz: {request.thema}", "fragen": [{{"nummer": 1, "frage": "...", "optionen": ["A) ...", "B) ...", "C) ...", "D) ..."], "richtig": "A", "erklaerung": "..."}}]}}""",

            "raetsel": f"""Erstelle ein Kreuzworträtsel für Deutsch RS+ Klasse {request.klassenstufe} zum Thema "{request.thema}" (Niveau: {niveau_name}).

Das Rätsel soll 8-10 Begriffe enthalten.

Format als JSON:
{{"titel": "Kreuzworträtsel: {request.thema}", "begriffe": [{{"wort": "...", "hinweis": "...", "richtung": "waagerecht/senkrecht", "nummer": 1}}], "loesung": ["Liste aller Lösungswörter"]}}""",

            "zuordnung": f"""Erstelle eine Zuordnungsübung für Deutsch RS+ Klasse {request.klassenstufe} zum Thema "{request.thema}" (Niveau: {niveau_name}).

Die Übung soll 8 Paare zum Zuordnen enthalten (z.B. Begriff → Definition, Beispiel → Regel).

Format als JSON:
{{"titel": "Zuordnung: {request.thema}", "anleitung": "Ordne die passenden Paare zu.", "paare": [{{"links": "...", "rechts": "..."}}], "tipp": "Ein hilfreicher Hinweis"}}""",

            "lueckentext": f"""Erstelle einen Lückentext für Deutsch RS+ Klasse {request.klassenstufe} zum Thema "{request.thema}" (Niveau: {niveau_name}).

Der Text soll 8-10 Lücken enthalten.

Format als JSON:
{{"titel": "Lückentext: {request.thema}", "text": "Der Text mit ___(1)___ Lücken ___(2)___ markiert...", "luecken": [{{"nummer": 1, "loesung": "...", "hinweis": "..."}}], "woerter_box": ["Liste der einzusetzenden Wörter (gemischt)"]}}"""
        }
        
        prompt = material_prompts.get(request.material_typ, material_prompts["arbeitsblatt"])
        
        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"material-{user_id}-{uuid.uuid4()}",
            system_message="""Du bist ein erfahrener Deutschlehrer an einer Realschule plus. 
Du erstellst kreative, schülergerechte Unterrichtsmaterialien.
Antworte IMMER nur mit validem JSON, ohne Erklärungen."""
        ).with_model("gemini", "gemini-3-flash-preview")
        
        response = await asyncio.wait_for(
            chat.send_message(UserMessage(text=prompt)),
            timeout=45.0
        )
        
        response_text = response.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        material = json_lib.loads(response_text.strip())
        
        return {
            "typ": request.material_typ,
            "niveau": niveau_name,
            "klassenstufe": request.klassenstufe,
            "thema": request.thema,
            "material": material
        }
        
    except json_lib.JSONDecodeError as e:
        logger.error(f"JSON Parse error: {e}")
        raise HTTPException(status_code=500, detail="Fehler beim Parsen der KI-Antwort")
    except asyncio.TimeoutError:
        raise HTTPException(status_code=504, detail="KI-Anfrage Timeout")
    except Exception as e:
        logger.error(f"Material generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============== STUNDEN-MATERIAL (AUFGABEN- UND LÖSUNGSBLÄTTER) ==============

class StundenMaterialRequest(BaseModel):
    stunde_titel: str
    stunde_inhalt: str
    stunde_aufgaben: list  # Liste der Aufgaben aus der Stunde
    stunde_lernziel: Optional[str] = None
    klassenstufe: str
    niveau: str
    thema: str

@router.post("/stunde/material/generieren")
async def generiere_stunden_material(
    request: StundenMaterialRequest,
    user_id: str = Depends(get_current_user)
):
    """Generiert Aufgaben- und Lösungsblätter für eine spezifische Unterrichtsstunde"""
    import json as json_lib
    
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        emergent_key = os.environ.get("EMERGENT_LLM_KEY", "")
        if not emergent_key:
            raise HTTPException(status_code=500, detail="KI-Service nicht konfiguriert")
        
        niveau_name = {"G": "grundlegend", "M": "mittel", "E": "erweitert"}.get(request.niveau, "mittel")
        
        # Aufgaben als String formatieren
        aufgaben_text = "\n".join([f"- {a}" for a in request.stunde_aufgaben]) if request.stunde_aufgaben else "Keine spezifischen Aufgaben angegeben"
        
        prompt = f"""Erstelle ein detailliertes Aufgabenblatt mit Lösungen für folgende Unterrichtsstunde:

STUNDENDETAILS:
- Titel: {request.stunde_titel}
- Thema: {request.thema}
- Klassenstufe: {request.klassenstufe}
- Niveau: {niveau_name}
- Lernziel: {request.stunde_lernziel or 'Nicht angegeben'}
- Inhalt: {request.stunde_inhalt}
- Geplante Aufgaben/Aktivitäten:
{aufgaben_text}

Erstelle basierend auf diesen Stundeninhalten ein KOMPLETTES Aufgabenblatt mit Lösungen.

Format als JSON:
{{
    "aufgabenblatt": {{
        "titel": "Aufgabenblatt: {request.stunde_titel}",
        "klassenstufe": "{request.klassenstufe}",
        "lernziel": "Das konkrete Lernziel",
        "aufgaben": [
            {{
                "nummer": 1,
                "titel": "Aufgabentitel",
                "aufgabenstellung": "Detaillierte Aufgabenstellung mit klaren Anweisungen",
                "punkte": 4,
                "material": "Optionaler Arbeitstext oder Tabelle falls nötig",
                "platz_fuer_antwort": true
            }}
        ],
        "gesamtpunkte": 20
    }},
    "loesungsblatt": {{
        "titel": "Lösungen: {request.stunde_titel}",
        "loesungen": [
            {{
                "nummer": 1,
                "loesung": "Die vollständige Musterlösung",
                "hinweise": "Optionale Korrekturhinweise für Lehrer"
            }}
        ]
    }}
}}

WICHTIG:
- Erstelle 4-6 abwechslungsreiche Aufgaben, die DIREKT zu den geplanten Aktivitäten passen
- Die Aufgaben sollen zum Niveau "{niveau_name}" passen
- Jede Aufgabe braucht eine klare, schülergerechte Aufgabenstellung
- Die Lösungen müssen vollständig und für Lehrer korrigierfreundlich sein
- Bei Textaufgaben: Füge einen passenden Arbeitstext als "material" hinzu
- Nur valides JSON zurückgeben"""

        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"stunden-material-{user_id}-{uuid.uuid4()}",
            system_message="""Du bist ein erfahrener Deutschlehrer an einer Realschule plus. 
Du erstellst professionelle Arbeitsblätter mit klaren Aufgabenstellungen und vollständigen Lösungen.
Deine Aufgaben sind abwechslungsreich und auf das Niveau der Schüler abgestimmt.
Antworte IMMER nur mit validem JSON, ohne Erklärungen."""
        ).with_model("gemini", "gemini-3-flash-preview")
        
        response = await asyncio.wait_for(
            chat.send_message(UserMessage(text=prompt)),
            timeout=60.0
        )
        
        response_text = response.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        material = json_lib.loads(response_text.strip())
        
        return {
            "stunde_titel": request.stunde_titel,
            "klassenstufe": request.klassenstufe,
            "niveau": niveau_name,
            "aufgabenblatt": material.get("aufgabenblatt", {}),
            "loesungsblatt": material.get("loesungsblatt", {})
        }
        
    except json_lib.JSONDecodeError as e:
        logger.error(f"JSON Parse error: {e}")
        raise HTTPException(status_code=500, detail="Fehler beim Parsen der KI-Antwort")
    except asyncio.TimeoutError:
        raise HTTPException(status_code=504, detail="KI-Anfrage Timeout")
    except Exception as e:
        logger.error(f"Stunden-Material generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/stunde/material/export/word")
async def export_stunden_material_to_word(
    aufgabenblatt: dict,
    loesungsblatt: dict,
    user_id: str = Depends(get_current_user)
):
    """Exportiert Aufgaben- und Lösungsblatt als separate Word-Dokumente in einem ZIP"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import zipfile
    
    def create_aufgabenblatt(data):
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Titel
        title = doc.add_heading(data.get("titel", "Aufgabenblatt"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadaten
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_text = f"Klassenstufe {data.get('klassenstufe', '')} • Name: _________________ • Datum: _________"
        meta.add_run(meta_text).italic = True
        
        # Lernziel
        if data.get("lernziel"):
            doc.add_paragraph()
            lz = doc.add_paragraph()
            lz.add_run("Lernziel: ").bold = True
            lz.add_run(data["lernziel"])
        
        doc.add_paragraph()
        
        # Aufgaben
        for aufgabe in data.get("aufgaben", []):
            # Aufgabenkopf
            heading = doc.add_paragraph()
            heading.add_run(f"Aufgabe {aufgabe.get('nummer', '')}: {aufgabe.get('titel', '')}").bold = True
            heading.add_run(f" ({aufgabe.get('punkte', 0)} Punkte)")
            
            # Aufgabenstellung
            if aufgabe.get("aufgabenstellung"):
                doc.add_paragraph(aufgabe["aufgabenstellung"])
            
            # Material/Text
            if aufgabe.get("material"):
                doc.add_paragraph()
                mat_para = doc.add_paragraph()
                mat_para.add_run("Material/Text:").bold = True
                doc.add_paragraph(aufgabe["material"])
            
            # Platz für Antwort
            if aufgabe.get("platz_fuer_antwort", True):
                doc.add_paragraph()
                doc.add_paragraph("Deine Antwort:")
                for _ in range(4):
                    doc.add_paragraph("_" * 70)
            
            doc.add_paragraph()
        
        # Gesamtpunkte
        total = doc.add_paragraph()
        total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total.add_run(f"Gesamtpunkte: ___ / {data.get('gesamtpunkte', 20)}").bold = True
        
        return doc
    
    def create_loesungsblatt(data):
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Titel
        title = doc.add_heading(data.get("titel", "Lösungsblatt"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Lösungen
        for loesung in data.get("loesungen", []):
            heading = doc.add_paragraph()
            heading.add_run(f"Lösung {loesung.get('nummer', '')}:").bold = True
            
            doc.add_paragraph(loesung.get("loesung", ""))
            
            if loesung.get("hinweise"):
                hint = doc.add_paragraph()
                hint.add_run("Korrekturhinweis: ").italic = True
                hint.add_run(loesung["hinweise"]).italic = True
            
            doc.add_paragraph()
        
        return doc
    
    try:
        # Erstelle beide Dokumente
        aufgaben_doc = create_aufgabenblatt(aufgabenblatt)
        loesung_doc = create_loesungsblatt(loesungsblatt)
        
        # Speichere in BytesIO
        aufgaben_buffer = BytesIO()
        loesung_buffer = BytesIO()
        aufgaben_doc.save(aufgaben_buffer)
        loesung_doc.save(loesung_buffer)
        aufgaben_buffer.seek(0)
        loesung_buffer.seek(0)
        
        # Erstelle ZIP
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            titel = aufgabenblatt.get("titel", "Material").replace(":", "-").replace("/", "-")
            zipf.writestr(f"{titel}_Aufgabenblatt.docx", aufgaben_buffer.read())
            zipf.writestr(f"{titel}_Loesungsblatt.docx", loesung_buffer.read())
        
        zip_buffer.seek(0)
        
        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename=Stundenmaterial.zip"}
        )
        
    except Exception as e:
        logger.error(f"Word export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============== GESPEICHERTE UNTERRICHTSREIHEN ==============

@router.get("/unterrichtsreihen")
async def get_saved_unterrichtsreihen(user_id: str = Depends(get_current_user)):
    """Gibt alle gespeicherten Unterrichtsreihen des Nutzers zurück"""
    db = get_db()
    cursor = db.unterrichtsreihen.find({"user_id": user_id})
    reihen = []
    async for doc in cursor:
        reihen.append({
            "id": str(doc["_id"]),
            "klassenstufe": doc.get("klassenstufe"),
            "kompetenzbereich": doc.get("kompetenzbereich"),
            "thema_id": doc.get("thema_id"),
            "niveau": doc.get("niveau"),
            "unterrichtsreihe": doc.get("unterrichtsreihe"),
            "created_at": doc.get("created_at")
        })
    return {"unterrichtsreihen": reihen}


@router.delete("/unterrichtsreihe/{reihe_id}")
async def delete_unterrichtsreihe(reihe_id: str, user_id: str = Depends(get_current_user)):
    """Löscht eine gespeicherte Unterrichtsreihe"""
    from bson import ObjectId
    db = get_db()
    try:
        result = await db.unterrichtsreihen.delete_one({
            "_id": ObjectId(reihe_id),
            "user_id": user_id
        })
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Unterrichtsreihe nicht gefunden")
        return {"success": True, "message": "Unterrichtsreihe gelöscht"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============== WORD EXPORT ==============

from fastapi.responses import StreamingResponse
from io import BytesIO
from pydantic import BaseModel as PydanticBaseModel

class WordExportRequest(PydanticBaseModel):
    material_typ: str
    titel: str
    inhalt: dict

@router.post("/material/export/word")
async def export_material_to_word(
    request: WordExportRequest,
    user_id: str = Depends(get_current_user)
):
    """
    Exportiert generiertes Material als Word-Dokument (.docx)
    Unterstützt: arbeitsblatt, quiz, raetsel, zuordnung
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    
    material_typ = request.material_typ
    titel = request.titel
    inhalt = request.inhalt
    
    doc = Document()
    
    # Seitenränder setzen
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Titel
    title = doc.add_heading(titel, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadaten
    if inhalt.get("klassenstufe") or inhalt.get("niveau"):
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_text = []
        if inhalt.get("klassenstufe"):
            meta_text.append(f"Klassenstufe {inhalt['klassenstufe']}")
        if inhalt.get("niveau"):
            niveau_map = {"G": "Grundniveau", "M": "Mittleres Niveau", "E": "Erweitertes Niveau"}
            meta_text.append(niveau_map.get(inhalt["niveau"], inhalt["niveau"]))
        meta.add_run(" • ".join(meta_text)).italic = True
    
    doc.add_paragraph()  # Leerzeile
    
    if material_typ == "arbeitsblatt":
        # Arbeitsblatt-Struktur
        aufgaben = inhalt.get("aufgaben", [])
        for i, aufgabe in enumerate(aufgaben, 1):
            # Aufgabentitel
            heading = doc.add_heading(f"Aufgabe {i}: {aufgabe.get('titel', '')}", level=2)
            
            # Aufgabenstellung
            if aufgabe.get("aufgabenstellung"):
                p = doc.add_paragraph(aufgabe["aufgabenstellung"])
            
            # Material/Text falls vorhanden
            if aufgabe.get("material"):
                doc.add_paragraph()
                material_para = doc.add_paragraph()
                material_para.add_run("Material: ").bold = True
                material_para.add_run(aufgabe["material"])
            
            # Platz für Antworten
            doc.add_paragraph()
            answer_lines = doc.add_paragraph("Antwort:")
            answer_lines.add_run("\n" + "_" * 60)
            answer_lines.add_run("\n" + "_" * 60)
            answer_lines.add_run("\n" + "_" * 60)
            
            doc.add_paragraph()  # Abstand zwischen Aufgaben
    
    elif material_typ == "quiz":
        # Quiz-Struktur
        fragen = inhalt.get("fragen", [])
        for i, frage in enumerate(fragen, 1):
            # Frage
            q_para = doc.add_paragraph()
            q_para.add_run(f"{i}. {frage.get('frage', '')}").bold = True
            
            # Antwortmöglichkeiten
            optionen = frage.get("optionen", [])
            for j, option in enumerate(optionen):
                # Prüfen ob Option bereits mit Buchstabe beginnt (z.B. "A) ...")
                option_text = str(option)
                if len(option_text) > 2 and option_text[1] in ')':
                    # Option hat bereits Buchstabe, nicht nochmal hinzufügen
                    doc.add_paragraph(f"   {option_text}")
                else:
                    buchstabe = chr(65 + j)  # A, B, C, D...
                    doc.add_paragraph(f"   {buchstabe}) {option_text}")
            
            doc.add_paragraph()  # Abstand
        
        # Lösungen am Ende
        doc.add_page_break()
        doc.add_heading("Lösungen", level=1)
        for i, frage in enumerate(fragen, 1):
            loesung = frage.get("loesung", "")
            doc.add_paragraph(f"{i}. {loesung}")
    
    elif material_typ == "raetsel":
        # Echtes Kreuzworträtsel mit sich kreuzenden Wörtern
        from docx.shared import Pt, Cm, RGBColor, Twips
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        doc.add_heading("Kreuzworträtsel", level=1)
        doc.add_paragraph()
        
        # Hole alle Begriffe
        begriffe = inhalt.get("begriffe", [])
        waagerecht = inhalt.get("waagerecht", [])
        senkrecht = inhalt.get("senkrecht", [])
        
        # Falls begriffe-Format verwendet wird, konvertiere es
        if begriffe and not waagerecht:
            waagerecht = [b for b in begriffe if b.get("richtung") == "waagerecht"]
            senkrecht = [b for b in begriffe if b.get("richtung") == "senkrecht"]
        
        # Sammle alle Wörter mit Infos
        h_words = []
        for item in waagerecht:
            wort = item.get("loesung", item.get("wort", "")).upper().strip()
            if wort:
                h_words.append({
                    "wort": wort, 
                    "nummer": item.get("nummer", len(h_words)+1),
                    "hinweis": item.get("frage", item.get("hinweis", ""))
                })
        
        v_words = []
        for item in senkrecht:
            wort = item.get("loesung", item.get("wort", "")).upper().strip()
            if wort:
                v_words.append({
                    "wort": wort, 
                    "nummer": item.get("nummer", len(v_words)+1),
                    "hinweis": item.get("frage", item.get("hinweis", ""))
                })
        
        # Kreuzworträtsel-Grid erstellen
        # Einfacher Algorithmus: Platziere Wörter so, dass sie sich überschneiden
        grid_size = 15
        grid = [[None for _ in range(grid_size)] for _ in range(grid_size)]
        placed_words = []
        nummer_positionen = []  # Speichert (row, col, nummer)
        
        # Erstes horizontales Wort in der Mitte platzieren
        if h_words:
            first_word = h_words[0]
            start_row = grid_size // 3
            start_col = (grid_size - len(first_word["wort"])) // 2
            for i, char in enumerate(first_word["wort"]):
                grid[start_row][start_col + i] = char
            placed_words.append({
                "wort": first_word["wort"],
                "row": start_row,
                "col": start_col,
                "dir": "H",
                "nummer": first_word["nummer"],
                "hinweis": first_word["hinweis"]
            })
            nummer_positionen.append((start_row, start_col, first_word["nummer"]))
        
        # Versuche vertikale Wörter zu platzieren (kreuzend)
        for v_word in v_words:
            wort = v_word["wort"]
            placed = False
            
            # Suche nach Überschneidungspunkt mit platzierten Wörtern
            for pw in placed_words:
                if pw["dir"] == "H":
                    # Suche gemeinsamen Buchstaben
                    for i, char in enumerate(wort):
                        if char in pw["wort"]:
                            # Gefunden! Platziere vertikal
                            h_idx = pw["wort"].index(char)
                            cross_col = pw["col"] + h_idx
                            cross_row = pw["row"]
                            start_row = cross_row - i
                            
                            # Prüfe ob Platzierung möglich
                            if start_row >= 0 and start_row + len(wort) < grid_size:
                                can_place = True
                                for j, c in enumerate(wort):
                                    r = start_row + j
                                    if grid[r][cross_col] is not None and grid[r][cross_col] != c:
                                        can_place = False
                                        break
                                
                                if can_place:
                                    for j, c in enumerate(wort):
                                        grid[start_row + j][cross_col] = c
                                    placed_words.append({
                                        "wort": wort,
                                        "row": start_row,
                                        "col": cross_col,
                                        "dir": "V",
                                        "nummer": v_word["nummer"],
                                        "hinweis": v_word["hinweis"]
                                    })
                                    nummer_positionen.append((start_row, cross_col, v_word["nummer"]))
                                    placed = True
                                    break
                if placed:
                    break
            
            # Fallback: Platziere separat
            if not placed and len(placed_words) < 8:
                col = 1 + len([p for p in placed_words if p["dir"] == "V"]) * 3
                row = 1
                if col < grid_size - 1:
                    for j, c in enumerate(wort):
                        if row + j < grid_size:
                            grid[row + j][col] = c
                    placed_words.append({
                        "wort": wort,
                        "row": row,
                        "col": col,
                        "dir": "V",
                        "nummer": v_word["nummer"],
                        "hinweis": v_word["hinweis"]
                    })
                    nummer_positionen.append((row, col, v_word["nummer"]))
        
        # Weitere horizontale Wörter platzieren
        for h_word in h_words[1:]:
            wort = h_word["wort"]
            placed = False
            
            for pw in placed_words:
                if pw["dir"] == "V":
                    for i, char in enumerate(wort):
                        if char in pw["wort"]:
                            v_idx = pw["wort"].index(char)
                            cross_row = pw["row"] + v_idx
                            cross_col = pw["col"]
                            start_col = cross_col - i
                            
                            if start_col >= 0 and start_col + len(wort) < grid_size:
                                can_place = True
                                for j, c in enumerate(wort):
                                    col = start_col + j
                                    if grid[cross_row][col] is not None and grid[cross_row][col] != c:
                                        can_place = False
                                        break
                                
                                if can_place:
                                    for j, c in enumerate(wort):
                                        grid[cross_row][start_col + j] = c
                                    placed_words.append({
                                        "wort": wort,
                                        "row": cross_row,
                                        "col": start_col,
                                        "dir": "H",
                                        "nummer": h_word["nummer"],
                                        "hinweis": h_word["hinweis"]
                                    })
                                    nummer_positionen.append((cross_row, start_col, h_word["nummer"]))
                                    placed = True
                                    break
                if placed:
                    break
        
        # Finde die tatsächlich verwendeten Grenzen des Gitters
        min_row, max_row = grid_size, 0
        min_col, max_col = grid_size, 0
        for r in range(grid_size):
            for c in range(grid_size):
                if grid[r][c] is not None:
                    min_row = min(min_row, r)
                    max_row = max(max_row, r)
                    min_col = min(min_col, c)
                    max_col = max(max_col, c)
        
        # Etwas Rand hinzufügen
        min_row = max(0, min_row - 1)
        min_col = max(0, min_col - 1)
        max_row = min(grid_size - 1, max_row + 1)
        max_col = min(grid_size - 1, max_col + 1)
        
        # Tabelle erstellen
        rows = max_row - min_row + 1
        cols = max_col - min_col + 1
        
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Zellen formatieren
        cell_size = Cm(0.7)
        for r_idx, row in enumerate(table.rows):
            row.height = cell_size
            for c_idx, cell in enumerate(row.cells):
                cell.width = cell_size
                
                grid_r = min_row + r_idx
                grid_c = min_col + c_idx
                
                # Vertikale Zentrierung
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if grid[grid_r][grid_c] is not None:
                    # Zelle mit Buchstabe - weißer Hintergrund, Rahmen
                    # Prüfe ob hier eine Nummer hingehört
                    nummer = None
                    for np in nummer_positionen:
                        if np[0] == grid_r and np[1] == grid_c:
                            nummer = np[2]
                            break
                    
                    if nummer:
                        # Kleine Nummer oben links
                        run = para.add_run(str(nummer))
                        run.font.size = Pt(6)
                        run.font.bold = False
                    
                    # Setze weißen Hintergrund
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFFFF"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                else:
                    # Leere Zelle - schwarzer Hintergrund
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="000000"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
        
        # Rahmen für alle Zellen
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="single" w:sz="4" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:color="000000"/>'
            '<w:insideH w:val="single" w:sz="4" w:color="000000"/>'
            '<w:insideV w:val="single" w:sz="4" w:color="000000"/>'
            '</w:tblBorders>'
        )
        tblPr.append(tblBorders)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Hinweise (Fragen)
        h_placed = [p for p in placed_words if p["dir"] == "H"]
        v_placed = [p for p in placed_words if p["dir"] == "V"]
        
        if h_placed:
            doc.add_heading("Waagerecht →", level=2)
            for item in sorted(h_placed, key=lambda x: x["nummer"]):
                doc.add_paragraph(f"{item['nummer']}. {item['hinweis']}")
        
        doc.add_paragraph()
        
        if v_placed:
            doc.add_heading("Senkrecht ↓", level=2)
            for item in sorted(v_placed, key=lambda x: x["nummer"]):
                doc.add_paragraph(f"{item['nummer']}. {item['hinweis']}")
        
        # Lösungsseite
        doc.add_page_break()
        doc.add_heading("Lösungen", level=1)
        
        if h_placed:
            doc.add_heading("Waagerecht", level=2)
            for item in sorted(h_placed, key=lambda x: x["nummer"]):
                doc.add_paragraph(f"{item['nummer']}. {item['wort']}")
        
        if v_placed:
            doc.add_heading("Senkrecht", level=2)
            for item in sorted(v_placed, key=lambda x: x["nummer"]):
                doc.add_paragraph(f"{item['nummer']}. {item['wort']}")
    
    elif material_typ == "zuordnung":
        # Zuordnungsübung
        doc.add_heading("Zuordnungsübung", level=1)
        doc.add_paragraph("Verbinde die zusammengehörenden Begriffe:")
        doc.add_paragraph()
        
        paare = inhalt.get("paare", [])
        
        # Linke Spalte
        doc.add_heading("Begriffe", level=2)
        for i, paar in enumerate(paare, 1):
            doc.add_paragraph(f"{i}. {paar.get('links', '')}")
        
        doc.add_paragraph()
        
        # Rechte Spalte (gemischt)
        doc.add_heading("Zuordnungen (durcheinander)", level=2)
        import random
        rechts_liste = [p.get("rechts", "") for p in paare]
        random.shuffle(rechts_liste)
        for buchstabe, item in zip("ABCDEFGHIJKLMNOP", rechts_liste):
            doc.add_paragraph(f"{buchstabe}. {item}")
        
        # Lösungen
        doc.add_page_break()
        doc.add_heading("Lösungen", level=1)
        for i, paar in enumerate(paare, 1):
            doc.add_paragraph(f"{i}. {paar.get('links', '')} → {paar.get('rechts', '')}")
    
    elif material_typ == "lueckentext":
        # Lückentext formatieren
        text = inhalt.get("text", "")
        luecken = inhalt.get("luecken", [])
        woerter_box = inhalt.get("woerter_box", [])
        
        # Wörterbox anzeigen (wenn vorhanden)
        if woerter_box:
            doc.add_heading("Wörterbox", level=2)
            woerter_para = doc.add_paragraph()
            # Wörter in einer Box darstellen
            woerter_text = "   •   ".join(woerter_box)
            run = woerter_para.add_run(woerter_text)
            run.font.size = Pt(12)
            run.font.bold = True
            doc.add_paragraph()
        
        # Aufgabenstellung
        doc.add_heading("Aufgabe: Fülle die Lücken aus!", level=2)
        doc.add_paragraph()
        
        # Text mit Lücken
        # Ersetze \n\n durch echte Absätze
        paragraphs = text.replace("\\n\\n", "\n\n").replace("\\n", "\n").split("\n\n")
        
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                # Text mit normaler Schrift
                run = p.add_run(para_text.strip())
                run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Hinweise zu den Lücken (optional)
        if luecken:
            doc.add_heading("Hinweise", level=2)
            for luecke in luecken:
                nummer = luecke.get("nummer", "")
                hinweis = luecke.get("hinweis", "")
                if hinweis:
                    doc.add_paragraph(f"({nummer}) {hinweis}")
        
        # Lösungsseite
        doc.add_page_break()
        doc.add_heading("Lösungen", level=1)
        
        if luecken:
            for luecke in luecken:
                nummer = luecke.get("nummer", "")
                loesung = luecke.get("loesung", "")
                doc.add_paragraph(f"({nummer}) {loesung}")
        
        # Vollständiger Lösungstext
        doc.add_paragraph()
        doc.add_heading("Vollständiger Text", level=2)
        
        # Ersetze Lücken durch Lösungen
        loesung_text = text
        for luecke in luecken:
            nummer = luecke.get("nummer", "")
            loesung = luecke.get("loesung", "")
            loesung_text = loesung_text.replace(f"__({nummer})__", f"**{loesung}**")
        
        loesung_paragraphs = loesung_text.replace("\\n\\n", "\n\n").replace("\\n", "\n").split("\n\n")
        for para_text in loesung_paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                # Text parsen für fett markierte Wörter
                parts = para_text.strip().split("**")
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.size = Pt(11)
                    if i % 2 == 1:  # Ungerade Indizes sind die Lösungswörter
                        run.font.bold = True
                        run.font.underline = True
    
    else:
        # Generischer Export
        doc.add_paragraph(str(inhalt))
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"Erstellt mit PlanEd • {datetime.now().strftime('%d.%m.%Y')}").italic = True
    
    # In BytesIO speichern
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Dateiname generieren
    safe_titel = "".join(c for c in titel if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
    filename = f"{safe_titel}_{material_typ}.docx"
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ============== QR-CODE GENERATOR ==============

@router.post("/material/qrcode")
async def generate_qr_code(
    url: str,
    titel: str = "Material",
    user_id: str = Depends(get_current_user)
):
    """Generiert einen QR-Code für eine URL"""
    import qrcode
    from io import BytesIO
    
    # QR-Code erstellen
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)
    
    # Als Bild generieren
    img = qr.make_image(fill_color="black", back_color="white")
    
    # In BytesIO speichern
    img_stream = BytesIO()
    img.save(img_stream, format='PNG')
    img_stream.seek(0)
    
    return StreamingResponse(
        img_stream,
        media_type="image/png",
        headers={"Content-Disposition": f"attachment; filename=qrcode_{titel}.png"}
    )


# ============== LEARNINGAPPS INTEGRATION ==============

@router.get("/learningapps/vorlagen")
async def get_learningapps_templates(
    material_typ: str = "quiz",
    user_id: str = Depends(get_current_user)
):
    """
    Gibt LearningApps-Vorlagen für verschiedene Materialtypen zurück.
    Diese können direkt auf learningapps.org erstellt werden.
    """
    
    templates = {
        "quiz": {
            "name": "Multiple-Choice-Quiz",
            "url": "https://learningapps.org/create.php?new=24",
            "beschreibung": "Erstellen Sie ein Quiz mit Multiple-Choice-Fragen",
            "icon": "❓",
            "anleitung": [
                "1. Klicken Sie auf den Link",
                "2. Geben Sie Ihre Fragen und Antworten ein",
                "3. Markieren Sie die richtige Antwort",
                "4. Speichern und teilen Sie den Link"
            ]
        },
        "lueckentext": {
            "name": "Lückentext",
            "url": "https://learningapps.org/create.php?new=35",
            "beschreibung": "Interaktiver Lückentext zum Ausfüllen",
            "icon": "📝",
            "anleitung": [
                "1. Klicken Sie auf den Link",
                "2. Geben Sie Ihren Text ein",
                "3. Markieren Sie Wörter als Lücken mit *Wort*",
                "4. Speichern und teilen Sie den Link"
            ]
        },
        "zuordnung": {
            "name": "Paare zuordnen",
            "url": "https://learningapps.org/create.php?new=21",
            "beschreibung": "Begriffe einander zuordnen (Memory-Stil)",
            "icon": "🔗",
            "anleitung": [
                "1. Klicken Sie auf den Link",
                "2. Geben Sie zusammengehörende Paare ein",
                "3. Wählen Sie die Darstellung (Karten, Liste, etc.)",
                "4. Speichern und teilen Sie den Link"
            ]
        },
        "raetsel": {
            "name": "Kreuzworträtsel",
            "url": "https://learningapps.org/create.php?new=32",
            "beschreibung": "Interaktives Kreuzworträtsel online lösen",
            "icon": "🔤",
            "anleitung": [
                "1. Klicken Sie auf den Link",
                "2. Geben Sie Begriffe und Hinweise ein",
                "3. Das Rätsel wird automatisch generiert",
                "4. Speichern und teilen Sie den Link"
            ]
        },
        "millionenspiel": {
            "name": "Wer wird Millionär",
            "url": "https://learningapps.org/create.php?new=28",
            "beschreibung": "Quiz im Stil von 'Wer wird Millionär'",
            "icon": "💰",
            "anleitung": [
                "1. Klicken Sie auf den Link",
                "2. Erstellen Sie 15 Fragen mit steigender Schwierigkeit",
                "3. Fügen Sie Joker hinzu",
                "4. Speichern und teilen Sie den Link"
            ]
        },
        "gruppenzuordnung": {
            "name": "Gruppenzuordnung",
            "url": "https://learningapps.org/create.php?new=22",
            "beschreibung": "Begriffe in Kategorien einordnen",
            "icon": "📊",
            "anleitung": [
                "1. Klicken Sie auf den Link",
                "2. Erstellen Sie Kategorien (z.B. Nomen, Verben, Adjektive)",
                "3. Fügen Sie Begriffe hinzu",
                "4. Speichern und teilen Sie den Link"
            ]
        },
        "wortsuchraetsel": {
            "name": "Wortgitter / Suchsel",
            "url": "https://learningapps.org/create.php?new=34",
            "beschreibung": "Versteckte Wörter im Buchstabengitter finden",
            "icon": "🔍",
            "anleitung": [
                "1. Klicken Sie auf den Link",
                "2. Geben Sie die zu suchenden Wörter ein",
                "3. Das Gitter wird automatisch erstellt",
                "4. Speichern und teilen Sie den Link"
            ]
        },
        "timeline": {
            "name": "Zeitstrahl",
            "url": "https://learningapps.org/create.php?new=31",
            "beschreibung": "Ereignisse chronologisch einordnen",
            "icon": "📅",
            "anleitung": [
                "1. Klicken Sie auf den Link",
                "2. Fügen Sie Ereignisse mit Datum hinzu",
                "3. Schüler ordnen die Ereignisse ein",
                "4. Speichern und teilen Sie den Link"
            ]
        }
    }
    
    if material_typ in templates:
        return templates[material_typ]
    
    return {"alle_vorlagen": templates}


@router.post("/learningapps/suche")
async def search_learningapps(
    suchbegriff: str,
    kategorie: str = "Deutsch",
    user_id: str = Depends(get_current_user)
):
    """
    Generiert einen Suchlink für LearningApps.org
    """
    from urllib.parse import quote
    
    # Kategorien-Mapping für LearningApps
    kategorien = {
        "Deutsch": "5",
        "Mathematik": "2", 
        "Englisch": "6",
        "Geschichte": "10",
        "Biologie": "8",
        "Physik": "9",
        "Chemie": "16",
        "Geografie": "11",
        "Musik": "13",
        "Kunst": "14",
        "Sport": "15",
        "Religion": "17",
        "Politik": "12"
    }
    
    kat_id = kategorien.get(kategorie, "5")
    encoded_search = quote(suchbegriff)
    
    return {
        "suchbegriff": suchbegriff,
        "kategorie": kategorie,
        "such_url": f"https://learningapps.org/index.php?category={kat_id}&s={encoded_search}",
        "direkt_erstellen": f"https://learningapps.org/create.php",
        "empfohlene_apps": [
            {
                "name": f"LearningApps zu '{suchbegriff}' durchsuchen",
                "url": f"https://learningapps.org/index.php?category={kat_id}&s={encoded_search}",
                "beschreibung": f"Fertige interaktive Übungen zum Thema '{suchbegriff}'"
            }
        ]
    }


# ============== WORD MIT QR-CODE ==============

class WordExportWithQRRequest(PydanticBaseModel):
    material_typ: str
    titel: str
    inhalt: dict
    qr_url: str = None  # Optional: URL für QR-Code
    learningapps_typ: str = None  # Optional: LearningApps-Typ für Link

@router.post("/material/export/word-mit-qr")
async def export_material_with_qr(
    request: WordExportWithQRRequest,
    user_id: str = Depends(get_current_user)
):
    """
    Exportiert Material als Word-Dokument MIT QR-Code und LearningApps-Link
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import qrcode
    from io import BytesIO
    
    material_typ = request.material_typ
    titel = request.titel
    inhalt = request.inhalt
    
    doc = Document()
    
    # Seitenränder
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Titel
    title_para = doc.add_heading(titel, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # QR-Code und LearningApps-Box oben rechts
    if request.qr_url or request.learningapps_typ:
        # Info-Box erstellen
        info_table = doc.add_table(rows=1, cols=2)
        info_table.autofit = True
        
        left_cell = info_table.rows[0].cells[0]
        right_cell = info_table.rows[0].cells[1]
        
        # Linke Seite: LearningApps Info
        if request.learningapps_typ:
            templates = {
                "quiz": ("Multiple-Choice-Quiz", "https://learningapps.org/create.php?new=24"),
                "lueckentext": ("Lückentext", "https://learningapps.org/create.php?new=35"),
                "zuordnung": ("Paare zuordnen", "https://learningapps.org/create.php?new=21"),
                "raetsel": ("Kreuzworträtsel", "https://learningapps.org/create.php?new=32"),
            }
            
            if request.learningapps_typ in templates:
                name, url = templates[request.learningapps_typ]
                p = left_cell.paragraphs[0]
                p.add_run("🎮 Interaktiv online üben:\n").bold = True
                p.add_run(f"{name}\n")
                p.add_run(url).font.size = Pt(8)
        
        # Rechte Seite: QR-Code
        if request.qr_url:
            # QR-Code generieren
            qr = qrcode.QRCode(version=1, box_size=6, border=1)
            qr.add_data(request.qr_url)
            qr.make(fit=True)
            qr_img = qr.make_image(fill_color="black", back_color="white")
            
            # Als Bytes speichern
            qr_stream = BytesIO()
            qr_img.save(qr_stream, format='PNG')
            qr_stream.seek(0)
            
            # In Dokument einfügen
            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run()
            run.add_picture(qr_stream, width=Cm(2.5))
            p.add_run("\n📱 Scan für Mobile").font.size = Pt(8)
        
        doc.add_paragraph()
    
    # Metadaten
    if inhalt.get("klassenstufe") or inhalt.get("niveau"):
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_parts = []
        if inhalt.get("klassenstufe"):
            meta_parts.append(f"Klassenstufe {inhalt['klassenstufe']}")
        if inhalt.get("niveau"):
            niveau_map = {"G": "Grundniveau", "M": "Mittleres Niveau", "E": "Erweitertes Niveau"}
            meta_parts.append(niveau_map.get(inhalt["niveau"], inhalt["niveau"]))
        meta.add_run(" • ".join(meta_parts)).italic = True
    
    doc.add_paragraph()
    
    # Material-spezifischer Inhalt (vereinfacht - nur Arbeitsblatt als Beispiel)
    if material_typ == "arbeitsblatt":
        aufgaben = inhalt.get("aufgaben", [])
        for i, aufgabe in enumerate(aufgaben, 1):
            doc.add_heading(f"Aufgabe {i}: {aufgabe.get('titel', '')}", level=2)
            if aufgabe.get("aufgabenstellung"):
                doc.add_paragraph(aufgabe["aufgabenstellung"])
            if aufgabe.get("material"):
                p = doc.add_paragraph()
                p.add_run("Material: ").bold = True
                p.add_run(aufgabe["material"])
            doc.add_paragraph("Antwort:\n" + "_" * 60 + "\n" + "_" * 60)
            doc.add_paragraph()
    
    elif material_typ == "quiz":
        fragen = inhalt.get("fragen", [])
        for i, frage in enumerate(fragen, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {frage.get('frage', '')}").bold = True
            for j, option in enumerate(frage.get("optionen", [])):
                option_text = str(option)
                if len(option_text) > 2 and option_text[1] in ')':
                    doc.add_paragraph(f"   {option_text}")
                else:
                    doc.add_paragraph(f"   {chr(65+j)}) {option_text}")
            doc.add_paragraph()
    
    else:
        # Generischer Export
        doc.add_paragraph(str(inhalt))
    
    # Footer mit QR-Code Hinweis
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_text = f"Erstellt mit PlanEd • {datetime.now().strftime('%d.%m.%Y')}"
    if request.qr_url:
        footer_text += " • 📱 QR-Code für mobile Nutzung"
    footer.add_run(footer_text).italic = True
    
    # Speichern
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    safe_titel = "".join(c for c in titel if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
    filename = f"{safe_titel}_{material_typ}_qr.docx"
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

